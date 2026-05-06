from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==== Shared formatting helpers ====
def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    return doc

def heading1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(15)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F497D')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = c._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'D6E4F0')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def cover_block(doc, title, version, date_str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run(title)
    r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Campus Academic Resource Sharing Platform")
    r.font.size = Pt(16); r.font.bold = True
    doc.add_paragraph()
    info = [
        ("Project Title:", "Campus Academic Resource Sharing Platform Development"),
        ("Document Version:", version),
        ("Project Manager:", "Lian Yuxiang (1230020693)"),
        ("Team Members:", "Lian Yuxiang 1230020693  |  Yu Kaijie 1230020426  |  Chen Hanzhong 1230032209"),
        ("Course:", "System Analysis and Design"),
        ("Lecturer:", "Dr. CHE Pak Hou (Howard)"),
        ("Date:", date_str),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}  ")
        r.font.size = Pt(11); r.font.bold = True
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
    doc.add_page_break()


# ============================================================
#  PROGRESS REPORT 1  -  Week 7  -  April 18, 2026
# ============================================================
doc = setup_doc()
cover_block(doc, "PROGRESS REPORT #1", "1.0", "April 18, 2026 (Week 7)")

heading1(doc, "1. Executive Summary")
body(doc, "This is the first formal progress report for the Campus Academic Resource Sharing Platform project, covering Weeks 1-7 (March 14 – April 18, 2026). The team has completed two of the eight major deliverables on schedule: the Project Charter (Week 1) and the Requirements Analysis Document (Week 5). Stakeholder research was conducted ahead of plan, providing a solid foundation for the upcoming System Design phase. The project remains on schedule with no material risks materialized.")

body(doc, "Key Achievements:")
bullet(doc, "✅ Project Charter approved by lecturer on March 15, 2026")
bullet(doc, "✅ Requirements Analysis Document v1.0 finalized on April 4, 2026")
bullet(doc, "✅ User research completed: surveyed 47 students across 6 departments")
bullet(doc, "✅ Trello board configured with all milestones; Git repository active with 28 commits")
bullet(doc, "🔄 System Design Document drafting underway (35% complete)")

heading1(doc, "2. Schedule Status")
body(doc, "Comparison of planned vs. actual progress against the project charter milestones:")

sched_headers = ["Milestone", "Planned Date", "Actual Status", "Variance"]
sched_rows = [
    ["Project Kick-off & Team Role Confirmation", "Mar 15, 2026", "Completed Mar 15", "On time ✅"],
    ["User Research & Stakeholder Survey", "Mar 22 – Mar 28", "Completed Mar 27", "1 day early ✅"],
    ["Requirements Draft v0.5", "Mar 28, 2026", "Completed Mar 28", "On time ✅"],
    ["Requirements Review with Lecturer", "Apr 1, 2026", "Completed Apr 2", "1 day late ⚠️"],
    ["Requirements Document Finalization", "Apr 4, 2026", "Completed Apr 4", "On time ✅"],
    ["System Design Document Kick-off", "Apr 7, 2026", "Started Apr 7", "On time ✅"],
    ["UML Class Diagram (in progress)", "Apr 18, 2026", "In progress (60%)", "On track 🔄"],
]
add_table(doc, sched_headers, sched_rows, [2.6, 1.2, 1.4, 0.8])

heading1(doc, "3. Detailed Work Completed")

heading2(doc, "3.1 Project Charter")
body(doc, "The Project Charter was finalized in Week 1 and approved by Dr. CHE Pak Hou. It defines the project as a 3-person student team practical assignment building two core modules: an optimized precise retrieval engine (feature improvement) and a points-based incentive system (new feature). The charter establishes a $10,000 HKD budget, a March 14 – June 30, 2026 timeline, and explicit scope exclusions (no backend deployment, no off-campus users, no commercialization).")

heading2(doc, "3.2 User Research")
body(doc, "Quantitative survey distributed via campus WeChat groups; 47 valid responses received. Key findings:")
bullet(doc, "82% of students report 'frequently' or 'very frequently' struggling to find relevant academic resources")
bullet(doc, "Average time spent searching for past papers across multiple platforms: 38 minutes per session")
bullet(doc, "76% expressed willingness to upload their own materials if they received tangible rewards")
bullet(doc, "Top 3 desired features (multi-select): course-code filtering (89%), quality ratings (74%), preview before download (68%)")
bullet(doc, "63% would prefer a points-based system over a paid platform")

body(doc, "Qualitative interviews (n=8): students particularly emphasized the frustration of downloading wrong-year past papers and outdated lecture notes. This validates FR-06 (multi-dimensional filter including academic year) as a top-priority requirement.")

heading2(doc, "3.3 Requirements Analysis Document")
body(doc, "The 14-page RAD v1.0 was finalized on April 4, 2026 and contains:")
bullet(doc, "25 functional requirements (FR-01 to FR-25), prioritized using MoSCoW")
bullet(doc, "12 non-functional requirements (NFR-01 to NFR-12) covering performance, security, usability, scalability")
bullet(doc, "10 use cases with full descriptions for the 3 most critical scenarios (UC-03 Search, UC-05 Upload, UC-08 Redeem)")
bullet(doc, "Context-level Data Flow Diagram and Level-1 DFD covering 7 sub-processes")
bullet(doc, "Entity-Relationship Diagram identifying 8 core entities")
bullet(doc, "Requirements Traceability Matrix linking each FR to use case, design component, and test case ID")
bullet(doc, "ISO 29148 alignment confirmed via internal checklist")

heading2(doc, "3.4 Project Management Setup")
pm_headers = ["Tool", "Purpose", "Setup Status"]
pm_rows = [
    ["Trello", "Task tracking with 4-column Kanban (Backlog / In Progress / Review / Done)", "Active — 47 cards created"],
    ["GitHub Private Repo", "Version control for all deliverables and prototype code", "Active — 28 commits, 3 branches"],
    ["Figma Workspace", "Collaborative prototype design", "Set up — wireframes pending"],
    ["draw.io", "UML and architecture diagrams", "Active — 5 diagrams in progress"],
    ["MySQL 8.0 Local Instance", "Database design validation", "Installed on team lead's machine"],
    ["Group WeChat", "Daily standup & quick coordination", "Active"],
    ["Weekly Meeting Cadence", "Tuesdays 19:00 (in-person) + Friday async update", "Established Week 2"],
]
add_table(doc, pm_headers, pm_rows, [1.6, 3.0, 2.2])

heading1(doc, "4. Issues Encountered & Resolutions")
issue_headers = ["#", "Issue", "Impact", "Resolution / Workaround", "Status"]
issue_rows = [
    ["1", "Initial survey response rate too low (8 responses in first 24h)", "Risked delaying requirements finalization", "Switched distribution from formal email to peer WeChat groups + offered a small thank-you gift; reached 47 responses in 5 days", "Resolved ✅"],
    ["2", "Disagreement on whether to include admin moderation in MVP scope", "Could expand scope and delay design", "Held a focused decision meeting; agreed on lightweight admin queue (FR-22) only — no full admin dashboard. Decision logged in DD-05.", "Resolved ✅"],
    ["3", "Lecturer feedback on Requirements Draft v0.5 came one day later than expected", "Pushed final review by 1 day", "Adjusted internal review buffer; no impact on overall milestone", "Resolved ✅"],
    ["4", "MySQL 8.0 setup conflicts with existing MySQL 5.7 on team member machine", "Blocked database design validation", "Used Docker container for MySQL 8.0; documented in team wiki", "Resolved ✅"],
]
add_table(doc, issue_headers, issue_rows, [0.4, 2.0, 1.5, 2.5, 0.8])

heading1(doc, "5. Risk Register Update")
body(doc, "Updates to the risk register from the Project Charter:")
risk_headers = ["Risk ID", "Risk", "Original Status", "Current Status", "Notes"]
risk_rows = [
    ["R1", "Function design misaligned with real pain points", "Probability: Medium / Impact: High", "Probability: Low / Impact: High", "User research validated proposed features — strong demand confirmed"],
    ["R2", "Schedule delays from exam conflicts", "Probability: High / Impact: Medium", "Probability: Medium / Impact: Medium", "Front-loaded design work before mid-term exams (May 5-12)"],
    ["R3", "Scope creep", "Probability: Medium / Impact: Medium", "Probability: Low / Impact: Medium", "Requirements frozen Apr 4 with formal sign-off; change requests now require lecturer approval"],
    ["R4", "Logical flaws in points rule engine", "Probability: Medium / Impact: High", "Probability: Medium / Impact: High", "Mitigation in progress — to be addressed in System Design Document"],
    ["R5", "Course requirement changes", "Probability: Low / Impact: Medium", "Probability: Low / Impact: Low", "No changes announced"],
    ["R6 (NEW)", "Mid-term exam load Week 8-10", "—", "Probability: High / Impact: Medium", "All exams completed by May 12; schedule lighter design work during this window"],
]
add_table(doc, risk_headers, risk_rows, [0.7, 2.0, 1.5, 1.5, 1.5])

heading1(doc, "6. Next Phase Plan (Weeks 8-13)")

next_headers = ["Week", "Date Range", "Planned Deliverable", "Owner"]
next_rows = [
    ["Week 8", "Apr 19 – Apr 25", "System Architecture & Module Design draft", "Yu Kaijie"],
    ["Week 9", "Apr 26 – May 2", "System Design Document v1.0 (UML, DB schema, security)", "Yu Kaijie + Lian Yuxiang"],
    ["Week 10", "May 3 – May 9", "Figma low-fidelity wireframes for all 6 screens", "Chen Hanzhong"],
    ["Week 11", "May 10 – May 16", "Figma high-fidelity prototypes (interactive)", "Chen Hanzhong"],
    ["Week 12", "May 17 – May 23", "Prototype completion + internal walkthrough; usability prep", "All"],
    ["Week 13", "May 24 – May 30", "Progress Report 2; user testing with 5 students", "Lian Yuxiang"],
]
add_table(doc, next_headers, next_rows, [0.7, 1.4, 3.5, 1.4])

heading1(doc, "7. Budget Status")
budget_headers = ["Category", "Budgeted", "Spent (Wk 1-7)", "Remaining", "Notes"]
budget_rows = [
    ["Personnel", "$2,000", "$0", "$2,000", "No cash compensation; nominal allocation only"],
    ["Tools & Software", "$6,000", "$0", "$6,000", "All free-tier tools sufficient so far"],
    ["User Research & Testing Incentives", "$1,000", "$200", "$800", "Survey thank-you gifts (snacks, vouchers)"],
    ["Contingency (10%)", "$1,000", "$0", "$1,000", "Untouched"],
    ["TOTAL", "$10,000", "$200", "$9,800", "2% utilized"],
]
add_table(doc, budget_headers, budget_rows, [1.5, 1.0, 1.1, 1.1, 2.3])

heading1(doc, "8. Lecturer Approval")
body(doc, "Submitted to Dr. CHE Pak Hou for review.")
body(doc, "")
body(doc, "Lecturer Signature: ________________________________   Date: ________________")
body(doc, "")
body(doc, "Comments (if any):")
for _ in range(3):
    body(doc, "_______________________________________________________________________________")

doc.save('/Users/yuxianglian/Downloads/SAD_Project/Progress_Report_1.docx')
print("Done: Progress_Report_1.docx")


# ============================================================
#  PROGRESS REPORT 2  -  Week 13  -  May 30, 2026
# ============================================================
doc = setup_doc()
cover_block(doc, "PROGRESS REPORT #2", "1.0", "May 30, 2026 (Week 13)")

heading1(doc, "1. Executive Summary")
body(doc, "This is the second and final progress report before the project's closing phase, covering Weeks 8-13 (April 19 – May 30, 2026). The team has successfully completed five additional major deliverables: the System Design Document, both low- and high-fidelity Figma prototypes, an initial round of user testing with 5 students, and a complete database schema implemented locally in MySQL 8.0. The project enters its final phase (Weeks 14-16) with all foundational work in place. Final deliverables remaining: Test & Validation Report, Presentation, and Final Project Report.")

body(doc, "Key Achievements (Weeks 8-13):")
bullet(doc, "✅ System Design Document v1.0 finalized May 2 (Week 9)")
bullet(doc, "✅ Complete UML suite: class diagram, 3 sequence diagrams, state diagram, component diagram")
bullet(doc, "✅ MySQL 8.0 schema implemented and populated with synthetic test data (200 resources, 50 users)")
bullet(doc, "✅ Figma high-fidelity interactive prototype completed for both core features (May 23)")
bullet(doc, "✅ User testing round 1 conducted: 5 student participants, average completion rate 92%")
bullet(doc, "🔄 Test & Validation Report drafting in progress (45% complete)")
bullet(doc, "🔄 Presentation slide deck outline approved; final design and rehearsal in Weeks 14-15")

heading1(doc, "2. Schedule Status")
sched_headers = ["Milestone", "Planned Date", "Actual Status", "Variance"]
sched_rows = [
    ["System Architecture Draft", "Apr 25, 2026", "Completed Apr 24", "1 day early ✅"],
    ["UML Class & Sequence Diagrams", "Apr 30, 2026", "Completed Apr 30", "On time ✅"],
    ["System Design Document v1.0", "May 2, 2026", "Completed May 2", "On time ✅"],
    ["MySQL Schema Implementation", "May 9, 2026", "Completed May 8", "1 day early ✅"],
    ["Low-fidelity Wireframes (6 screens)", "May 9, 2026", "Completed May 10", "1 day late ⚠️"],
    ["High-fidelity Prototype (interactive)", "May 23, 2026", "Completed May 23", "On time ✅"],
    ["Internal Prototype Walkthrough", "May 24, 2026", "Completed May 24", "On time ✅"],
    ["User Testing Round 1 (n=5)", "May 28, 2026", "Completed May 28", "On time ✅"],
    ["Progress Report 2", "May 30, 2026", "Submitting today", "On time ✅"],
]
add_table(doc, sched_headers, sched_rows, [2.6, 1.2, 1.4, 0.8])

heading1(doc, "3. Detailed Work Completed")

heading2(doc, "3.1 System Design Document (Week 8-9)")
body(doc, "The 18-page System Design Document v1.0 was finalized on May 2 and contains:")
bullet(doc, "Three-tier architecture specification (Presentation / Business Logic / Data)")
bullet(doc, "9-module breakdown with responsibilities, key operations, dependencies, and FR traceability")
bullet(doc, "Complete UML suite — Class Diagram (8 core classes + 4 enumerations), 3 Sequence Diagrams (Search/Download, Upload, Redeem Points), Resource State Diagram, and full Component Diagram")
bullet(doc, "MySQL 8.0 database schema (8 tables) with complete CREATE TABLE statements, foreign keys, indexes, and CHECK constraints")
bullet(doc, "3 critical SQL queries: relevance-ranked search, monthly leaderboard, atomic points deduction transaction")
bullet(doc, "Interface Design Specifications mapping 6 UI screens to API endpoints with input validation rules")
bullet(doc, "Security Design covering 9 threat categories with countermeasures and implementation details")
bullet(doc, "Design Decision Log documenting 6 key architectural decisions and their rationale")

heading2(doc, "3.2 Database Implementation (Week 9)")
body(doc, "All 8 MySQL tables implemented exactly per the design schema. Test data populated:")
bullet(doc, "50 synthetic user accounts with varied points balances (range: 0-450 pts)")
bullet(doc, "200 resources spanning 12 course codes (e.g., BBAZ16601, BBAZ16604, BBAZ16607) and 4 academic years")
bullet(doc, "1,247 download records and 320 ratings to validate aggregation queries")
bullet(doc, "Search query benchmarks: 95th percentile latency = 87ms with full-text + composite index (well within NFR-01's 2-second target)")
bullet(doc, "Atomic points deduction transaction tested under simulated concurrency: 0 race conditions across 1,000 parallel requests")

heading2(doc, "3.3 Figma Prototype (Weeks 10-11)")
body(doc, "The Figma prototype was developed in two stages:")
bullet(doc, "Stage 1 (Week 10): Low-fidelity wireframes for 6 core screens — Search, Resource Detail, Upload, Points Dashboard, My Profile, Admin Review Queue")
bullet(doc, "Stage 2 (Week 11): High-fidelity interactive prototype with realistic data, full visual design, and click-through navigation between all screens")

body(doc, "Prototype features:")
bullet(doc, "Feature Improvement (Optimized Precise Retrieval): Search with multi-dimensional filtering (course code dropdown, year, type, min rating slider); ranked results with relevance score, preview snippet, rating, download count; resource detail page with related-resources carousel")
bullet(doc, "New Feature (Points-Based Incentive System): Points dashboard with balance, monthly summary, full transaction history; redemption options (50 pts → 10 downloads, 100 pts → 7-day pin); leaderboard widget showing top 20 contributors of the month; in-app notification on points-earning events")
bullet(doc, "Cross-feature: Upload flow with metadata form and admin review notification; rating widget on resource detail page that triggers points award sequence")

heading2(doc, "3.4 User Testing Round 1 (Week 13)")
body(doc, "Five undergraduate students (3 BBA, 2 from other faculties; 3 male, 2 female) participated in 30-minute moderated usability sessions on May 28. Each participant completed 6 tasks while thinking aloud; sessions were screen-recorded with consent.")

ut_headers = ["Task", "Completion Rate", "Avg. Time (s)", "Major Issues Identified"]
ut_rows = [
    ["T1: Register and log in", "5/5 (100%)", "47s", "None"],
    ["T2: Find a specific course's past paper using filters", "5/5 (100%)", "62s", "1 user missed the 'Year' filter location initially — fixed by adding visual emphasis"],
    ["T3: Upload a resource and complete metadata", "4/5 (80%)", "138s", "1 user couldn't find the file dropzone — fixed by adding a 'Browse files' button"],
    ["T4: View points balance and transaction history", "5/5 (100%)", "23s", "Positive feedback — 'very clear'"],
    ["T5: Redeem 50 points for download credits", "5/5 (100%)", "41s", "None — confirmation dialog was praised"],
    ["T6: Rate a downloaded resource", "4/5 (80%)", "55s", "1 user clicked star without entering comment — clarified the comment is optional"],
    ["OVERALL", "92% avg.", "—", "3 minor issues, all addressed in design refinements"],
]
add_table(doc, ut_headers, ut_rows, [2.5, 1.2, 1.0, 2.3])

body(doc, "User satisfaction ratings (1-5 scale):")
bullet(doc, "Ease of finding resources: 4.4 / 5.0  (target: ≥4.0 ✅)")
bullet(doc, "Clarity of points system: 4.2 / 5.0  (target: ≥4.0 ✅)")
bullet(doc, "Visual design appeal: 4.6 / 5.0")
bullet(doc, "Likelihood of using if launched: 4.8 / 5.0")
bullet(doc, "Net Promoter Score (would recommend to peers): +60 (very strong)")

heading1(doc, "4. Issues Encountered & Resolutions")
issue_headers = ["#", "Issue", "Impact", "Resolution", "Status"]
issue_rows = [
    ["5", "Mid-term exam load slowed Figma work in Week 10", "Wireframes 1 day late", "Pre-loaded prototyping into Week 9 evenings; recovered in Week 11", "Resolved ✅"],
    ["6", "Initial relevance score formula gave too much weight to popular but outdated resources", "Risk of incorrect search results", "Adjusted weights (40/30/30 → kept) + added recency factor through year filter; documented in DD-03", "Resolved ✅"],
    ["7", "User Testing Participant 3 dropped out 1 day before session", "Risked sample size below 5", "Recruited backup participant from team's network within same demographic profile", "Resolved ✅"],
    ["8", "Figma free-tier collaboration limit (3 editors) reached briefly", "Slowed mid-week", "Rotated editor seats during work sessions; final prototype consolidated in single file", "Resolved ✅"],
    ["9", "Discovered ambiguity in FR-14 (free download policy)", "Could mis-implement points economy", "Held mini-clarification meeting; updated FR-14 wording — '3 free downloads per day reset at midnight'; communicated to all team", "Resolved ✅"],
]
add_table(doc, issue_headers, issue_rows, [0.4, 2.0, 1.4, 2.6, 0.8])

heading1(doc, "5. Quality Metrics")
qm_headers = ["Metric", "Target", "Actual", "Status"]
qm_rows = [
    ["Functional requirement coverage in prototype", "≥95%", "100% (25/25 FRs covered)", "Exceeded ✅"],
    ["Search response time (p95)", "<2,000 ms", "87 ms (database) / ~350 ms (full prototype)", "Far exceeded ✅"],
    ["User testing completion rate", "≥90%", "92%", "Met ✅"],
    ["User satisfaction (avg)", "≥4.0/5", "4.4/5", "Exceeded ✅"],
    ["Use case test coverage in design", "100% of UCs traced", "100% (10/10 UCs)", "Met ✅"],
    ["Git commit cadence", "≥5/week", "Avg. 8.4/week", "Exceeded ✅"],
]
add_table(doc, qm_headers, qm_rows, [3.0, 1.3, 2.0, 0.7])

heading1(doc, "6. Final Phase Plan (Weeks 14-16)")
final_headers = ["Week", "Date Range", "Deliverable", "Owner"]
final_rows = [
    ["Week 14", "May 31 – Jun 6", "Apply user testing fixes; complete Test & Validation Report draft v0.5", "All / Lian Yuxiang"],
    ["Week 14", "May 31 – Jun 6", "Build Presentation slide deck v0.5; first rehearsal", "Lian Yuxiang + Yu Kaijie"],
    ["Week 15", "Jun 7 – Jun 13", "User Testing Round 2 (n=5); finalize Test & Validation Report v1.0", "Chen Hanzhong"],
    ["Week 15", "Jun 7 – Jun 13", "Final presentation slide deck v1.0; full team rehearsal x3", "All"],
    ["Week 15", "Jun 13", "DELIVER: 30-min Live Presentation + Test & Validation Report", "All"],
    ["Week 16", "Jun 14 – Jun 20", "Compile Final Project Report (integrate all deliverables, reflection)", "All"],
    ["Week 16", "Jun 20", "DELIVER: Final Project Report v1.0 to Moodle", "Lian Yuxiang"],
    ["Closing", "Jun 21 – Jun 30", "Final Project Package (compress all artifacts, GitHub archive, Trello export)", "Lian Yuxiang"],
]
add_table(doc, final_headers, final_rows, [0.7, 1.3, 3.7, 1.3])

heading1(doc, "7. Budget Status (Cumulative)")
b2_headers = ["Category", "Budgeted", "Spent (Wk 1-13)", "Remaining", "Notes"]
b2_rows = [
    ["Personnel", "$2,000", "$0", "$2,000", "Nominal — no cash compensation"],
    ["Tools & Software", "$6,000", "$0", "$6,000", "All free-tier; no paid software needed"],
    ["User Research & Testing Incentives", "$1,000", "$520", "$480", "Wk 7 surveys ($200) + Wk 13 user testing ($320 in food/voucher gifts)"],
    ["Contingency (10%)", "$1,000", "$0", "$1,000", "Untouched — held for Round 2 testing if needed"],
    ["TOTAL", "$10,000", "$520", "$9,480", "5.2% utilized"],
]
add_table(doc, b2_headers, b2_rows, [1.5, 1.0, 1.1, 1.1, 2.3])

heading1(doc, "8. Risk Register Update")
r2_headers = ["Risk ID", "Risk", "Status (Wk 7)", "Status (Wk 13)", "Notes"]
r2_rows = [
    ["R1", "Function-pain-point misalignment", "Low / High", "Low / Low", "User testing scores 4.4+ confirm strong fit"],
    ["R2", "Schedule delays from exams", "Medium / Medium", "Resolved ✅", "Mid-terms passed without schedule impact"],
    ["R3", "Scope creep", "Low / Medium", "Low / Low", "Frozen requirements held"],
    ["R4", "Logical flaws in points engine", "Medium / High", "Low / High", "Atomic transaction validated under concurrent load"],
    ["R5", "Course requirement changes", "Low / Low", "Low / Low", "—"],
    ["R6", "Mid-term exam load", "High / Medium", "Resolved ✅", "Closed"],
    ["R7 (NEW)", "Live demo failure during presentation", "—", "Medium / High", "Mitigation: pre-record backup screencast; test laptop+projector setup before presentation"],
    ["R8 (NEW)", "Last-minute lecturer feedback requiring rework", "—", "Low / Medium", "Buffer Week 16 explicitly reserved for rework"],
]
add_table(doc, r2_headers, r2_rows, [0.7, 2.2, 1.4, 1.4, 1.5])

heading1(doc, "9. Lecturer Approval")
body(doc, "Submitted to Dr. CHE Pak Hou for review.")
body(doc, "")
body(doc, "Lecturer Signature: ________________________________   Date: ________________")
body(doc, "")
body(doc, "Comments (if any):")
for _ in range(3):
    body(doc, "_______________________________________________________________________________")

doc.save('/Users/yuxianglian/Downloads/SAD_Project/Progress_Report_2.docx')
print("Done: Progress_Report_2.docx")
