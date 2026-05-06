"""
Presentation_Script.docx v3 — 配套新 11 张幻灯片的 15 分钟讲稿（3 人轮讲）
新结构：Title · Agenda · Pain · Two Features · Architecture · Live Demo
       · DFD · v2.0 AI Rec · Result & Impact · Q&A · Thank You
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

def heading1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(12); r.font.bold = True
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs: p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(15)
    return p

def script(doc, text):
    p = doc.add_paragraph(text)
    if p.runs: p.runs[0].font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]; c.text = h
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1F497D'); shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            c = row.cells[ci]; c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = c._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'D6E4F0'); shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# === Cover ===
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("PRESENTATION SCRIPT (v2)")
r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
doc.add_paragraph()
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("MUST Campus Academic Resource Sharing Platform")
r.font.size = Pt(16); r.font.bold = True
doc.add_paragraph()
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run("校园学术资源共享平台 · 配套 Final_Presentation.html 的逐页讲稿（15 分钟 · 3 人轮讲）")
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x65, 0x75, 0x8B)
doc.add_paragraph()
info = [
    ("Slides:", "11 张  ·  时长约 15 分钟（含 Q&A）"),
    ("Speakers:", "陈瀚中（开场+痛点+两大功能）/ 连宇翔（LIVE DEMO）/ 郁凯杰（架构+DFD+v2.0+Impact+致谢）"),
    ("Course:", "System Analysis and Design"),
    ("Lecturer:", "Dr. CHE Pak Hou (Howard)"),
    ("URL:", "https://signing-isle-printed-shapes.trycloudflare.com"),
    ("Demo Account:", "学号 1230000000  ·  密码 demo123  ·  100 积分立即可用"),
]
for label, value in info:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}  "); r.font.size = Pt(11); r.font.bold = True
    r2 = p.add_run(value); r2.font.size = Pt(11)
doc.add_page_break()

# === 时间轴 ===
heading1(doc, "演讲时间轴 (15 分钟 · 3 人轮讲)")
schedule_headers = ["#", "时段", "主题", "主讲", "时长"]
schedule_rows = [
    ["1",     "0:00 – 0:30",   "Title · 开场",                "陈瀚中", "0:30"],
    ["2",     "0:30 – 1:00",   "Agenda · 路线图",              "陈瀚中", "0:30"],
    ["3",     "1:00 – 2:30",   "1.0 痛点 + 用户画像",          "陈瀚中", "1:30"],
    ["4",     "2:30 – 4:00",   "2.0 两大功能：检索 + 积分",     "陈瀚中", "1:30"],
    ["5",     "4:00 – 5:00",   "2.x 技术架构 · 真实部署",       "郁凯杰", "1:00"],
    ["6",     "5:00 – 9:30",   "🎬 LIVE DEMO",                "连宇翔", "4:30"],
    ["7",     "9:30 – 11:00",  "2.3.4 DFD · 积分系统数据流",    "郁凯杰", "1:30"],
    ["8",     "11:00 – 12:30", "3.0 v2.0 AI 推荐（未来功能）",   "郁凯杰", "1:30"],
    ["9",     "12:30 – 13:30", "4.0 Result & Impact · 项目影响","郁凯杰", "1:00"],
    ["10",    "13:30 – 14:30", "Q&A · 问答",                  "全员",   "1:00"],
    ["11",    "14:30 – 15:00", "Thank You · 致谢",            "郁凯杰", "0:30"],
]
add_table(doc, schedule_headers, schedule_rows, [0.4, 1.4, 2.9, 1.0, 0.7])
body(doc, "")
body(doc, "讲稿规则：方括号 [ ] 内是动作 / 暗示性提示词，不要念出口；引号「」内是建议台词，可酌情自然化。")

doc.add_page_break()

# ============================================================
# Slide-by-slide
# ============================================================
heading1(doc, "逐页讲稿")

# --- Slide 1 ---
heading2(doc, "▼ Slide 1：Title  (0:00 – 0:30  ·  陈瀚中)")
script(doc, "「各位老师、各位同学，下午好。我是陈瀚中，旁边是连宇翔、郁凯杰，我们是这个项目的三人组。」")
script(doc, "「接下来 15 分钟，我们带大家看一个真实在线运行的网站 —— MUST Campus Academic Resource Sharing Platform，中文名校园学术资源共享平台。」")
script(doc, "「这不是一份 PPT 项目，演讲结束后它还会 24 小时在线。」")

# --- Slide 2 ---
heading2(doc, "▼ Slide 2：Agenda  (0:30 – 1:00  ·  陈瀚中)")
script(doc, "「先讲 1 分半钟我们看到的痛点，再讲 1 分半钟我们做的两大功能。」")
script(doc, "「然后凯杰用 1 分钟介绍系统怎么搭起来。中间最大的一段 —— 4 分半钟的 LIVE DEMO，由宇翔现场操作。」")
script(doc, "「之后凯杰会用 3 分钟讲数据流图和 v2.0 AI 推荐，再用 1 分钟收 Result & Impact。最后是 Q&A 和致谢。」")

# --- Slide 3 ---
heading2(doc, "▼ Slide 3：1.0 痛点 + 用户画像  (1:00 – 2:30  ·  陈瀚中)")
script(doc, "「先抛一个问题给大家：上一次你为了找一份历年试卷或学姐笔记，平均花了多少时间？」")
script(doc, "「我们做了 47 份问卷 + 8 个深访。82% 的同学告诉我们 —— 经常找不到自己想要的资源。这不是个别痛点，是普遍困扰。」")
script(doc, "「右边三张数据卡更具体：89% 想要按课程代码筛选，76% 在有奖励的前提下愿意主动分享，63% 偏好积分制而非付费订阅。」")
script(doc, "「为什么这么难找？因为资源散在课程群、网盘、付费平台、学长私聊里。每个地方搜索都很粗，下错年份的真题、买到过期资料是常事。」")
script(doc, "「这给了我们一个机会：做一个统一的、能精准筛选的、用积分激励同学分享的平台。」")

# --- Slide 4 ---
heading2(doc, "▼ Slide 4：2.0 两大功能  (2:30 – 4:00  ·  陈瀚中)")
script(doc, "「我们做了一个网站，让学生能 1 分钟找到资源、3 步上传分享、被下载就有积分回馈。」")
script(doc, "「核心是两件事，按课程要求做了 1 个改进 + 1 个新功能。」")
script(doc, "「左边是改进 —— 优化精准检索。课程代码、学年、资源类型、最低评分四个维度同时筛选。结果排序是综合公式：文本匹配 40% + 下载热度 30% + 评分 30%。实测响应不到 100 毫秒。」")
script(doc, "「右边是新功能 —— 积分激励体系。注册送 100 分，上传被审核通过 +10，资源被下载 +2，被好评 +1。下载扣 1 分，每天 3 次免费下载兜底。50 分换 100 次下载券，100 分可以把自己上传的资源置顶 7 天。」")
script(doc, "「这两个功能不是独立的 —— 形成飞轮：你上传越多分越多，搜索质量越高，社区越活。」")

# --- Slide 5 ---
heading2(doc, "▼ Slide 5：2.x 技术架构  (4:00 – 5:00  ·  郁凯杰)")
script(doc, "「谢谢瀚中。我是郁凯杰，我用 1 分钟讲系统怎么搭起来。」")
script(doc, "「上面这条 pipeline 是经典三层架构 —— 单文件 HTML 前端 / FastAPI Python 中间层 / MariaDB 数据库。MariaDB 是 MySQL 的开源分叉，SQL 完全通用。」")
script(doc, "「下面这条 pipeline 是部署：1GB 内存的云服务器，用 Docker 容器隔离 db 和 api 互不挤兑；Cloudflare 提供合法 HTTPS 证书；URL 24 小时在线，老师现在打开就能访问。」")
script(doc, "「关键一句：这个系统不是只在我们电脑上跑的 demo。下一秒交给宇翔做现场演示。」")

# --- Slide 6 ---
heading2(doc, "▼ Slide 6：LIVE DEMO  (5:00 – 9:30  ·  连宇翔)")
script(doc, "「谢谢凯杰。我是连宇翔，接下来 4 分半钟我带大家做完整演示。」")
script(doc, "「请看屏幕，网站现已打开。注意右下角的『后端在线』指示，说明系统真在跑。」")
script(doc, "")

heading3(doc, "▶ Step 1: 当场注册账号 (5:00 – 6:00)")
script(doc, "「先做个最关键的动作 —— 我们当场注册一个账号，证明系统真的能用。」")
script(doc, "「我现在输一个 8 位学号、一个密码、一个邮箱。」")
script(doc, "「点注册 —— 立刻 100 积分到账。这一秒数据已经写入 MariaDB 的 users 表。」")
script(doc, "[备用方案：直接使用预置演示账号 1230000000 / demo123 登录，同样 100 积分立即可用]")

heading3(doc, "▶ Step 2: 搜索演示 (6:00 – 6:45)")
script(doc, "「我用刚注册的账号搜索 —— 关键字 Forecasting，加上课程代码筛选。」")
script(doc, "「点搜索 —— 100 毫秒以内出结果。每条卡片显示文本匹配度、下载量、平均评分，对应 40/30/30 公式。」")

heading3(doc, "▶ Step 3: 真下载 + 积分扣减 (6:45 – 7:30)")
script(doc, "「点开第一份资源 —— 商务预测第 1 讲讲义，1MB 左右。」")
script(doc, "「点下载，弹出确认框：『要扣 1 分』。这是花积分前必须明确告知的 UX 原则。」")
script(doc, "「确认 —— 文件真的下到本地，余额从 100 变成 99。这一笔扣分走的是数据库原子事务，等下凯杰会用 DFD 把它讲清楚。」")

heading3(doc, "▶ Step 4: 评分 (7:30 – 8:00)")
script(doc, "「给这份资源打 5 星，提交。」")
script(doc, "「上传者立刻收到 +1 积分（4 星以上才奖励），形成激励循环。」")

heading3(doc, "▶ Step 5: 积分中心 (8:00 – 8:45)")
script(doc, "「切到积分中心 —— 这是项目的核心新功能。」")
script(doc, "「顶部三个数字：余额、本月获得、今日剩余免费下载。下方是完整流水，每一笔变动都有时间、原因、剩余余额，可审计。」")
script(doc, "「右边兑换：50 分换 100 次下载券、100 分置顶 7 天。我们演示兑换 —— 余额减 50，下载券 +100。这一切都在原子事务里完成。」")

heading3(doc, "▶ Step 6: 管理后台 (8:45 – 9:30)")
script(doc, "「最后秀一下后台 —— 我切到团队管理员账号登录。」")
script(doc, "[切到管理员账号 1230020693 · 密码自己记得]")
script(doc, "「左侧『管理』菜单 —— 顶部 6 张卡片是全平台统计：用户、资源、积分流水、积分总池。」")
script(doc, "「下面 4 个 tab：Users 是全体用户、Resources 含所有状态、Points Ledger 是 100 条最近流水、Review Queue 是待审核队列。点 Approve —— 上传者立刻 +10 分。」")
script(doc, "「演示到此结束，把话筒交给凯杰讲数据流图。」")

# --- Slide 7 ---
heading2(doc, "▼ Slide 7：DFD · 积分系统数据流  (9:30 – 11:00  ·  郁凯杰)")
script(doc, "「谢谢宇翔。刚才大家看到宇翔点了一下下载，余额从 100 变成 99。」")
script(doc, "「这一刀是怎么扣的？看左边这张 DFD —— 这是积分系统的数据流图，把整个『事件 → 引擎 → 存储 → 返回』讲清楚。」")
script(doc, "「右边四步对照：第一步，事件触发。无论是上传通过、被下载、被评分、用户消费，都进入同一个 Points Engine。」")
script(doc, "「第二步，原子事务。Engine 用一条带 WHERE 守卫的 UPDATE，把『读余额-改余额-写余额』的三步压在数据库锁内一次完成 —— 这就是不会出现负余额的根本原因。」")
script(doc, "「第三步，双写。同一事务里，新余额写到 Users 表，审计行写到 PointRecord 表 —— 两张表要么都成功，要么都回滚。」")
script(doc, "「第四步，返回。Engine 把新余额回给 UI，前端 toast 提示「-1」或「+10」。」")
script(doc, "「最下面这行是真实压测：1000 笔并发扣分，零 race condition，余额永不为负。这就是 SDD 第 3.4 节『WHERE-guard 模式』真实运行的证据。」")

# --- Slide 8 ---
heading2(doc, "▼ Slide 8：3.0 v2.0 AI 推荐（未来功能）  (11:00 – 12:30  ·  郁凯杰)")
script(doc, "「v1.0 跑稳之后，我们计划在 v2.0 加 AI 推荐 —— 把『主动搜索』升级成『被动发现』。」")
script(doc, "「左边 WHY：高频用户每次重新筛选成本太高；新生不知道该搜什么关键词，需要冷启动支持；学长学姐的隐性经验值得转成可复用建议。」")
script(doc, "「中间 USE CASE：读取用户历史 → 提取兴趣向量 → 按专业 + 课程 tag 筛选候选 → 混合模型（内容相似 + 协同过滤）排序 → 每张卡附推荐理由（同课程 / 高分 / 同专业收藏）。」")
script(doc, "「右边 DFD + CARE：因为是学习数据，所以隐私必须做到位 —— 不暴露任何其他用户的具体行为；用户可单条 dismiss 或整体 opt-out；负反馈进入下一轮训练。占位克制，不抢现有搜索流的注意力。」")
script(doc, "「v2.0 不会取代 v1.0 的搜索，而是叠加在它上面。」")

# --- Slide 9 ---
heading2(doc, "▼ Slide 9：4.0 Result & Impact  (12:30 – 13:30  ·  郁凯杰)")
script(doc, "「最后 1 分钟讲项目影响。报告里写了两条线 —— 社会和财务。」")
script(doc, "「社会影响：建立校园学术互助文化；低质资源被评分淘汰；贡献者首次获得显式认可；课程热度数据可反馈给学生组织。」")
script(doc, "「财务影响：v1.0 不商业化，价值在节省时间；避免重复购买低质二手资料；积分未来可对接印店或图书馆兑换；商业化空间留给 v2.0。」")
script(doc, "「一句话收尾：一个真实部署、能用的系统，比 100 页文档项目说的更多。」")

# --- Slide 10 ---
heading2(doc, "▼ Slide 10：Q&A  (13:30 – 14:30  ·  全员)")
script(doc, "「OK，到 Q&A 环节。我们 3 个人都在台上 —— 技术问题凯杰熟，演示问题宇翔熟，整体规划瀚中接。请问大家有什么问题？」")
script(doc, "[预留约 60 秒 · 准备 3 个高概率问题：见后文]")

# --- Slide 11 ---
heading2(doc, "▼ Slide 11：Thank You  (14:30 – 15:00  ·  郁凯杰)")
script(doc, "「再次感谢老师和同学聆听。」")
script(doc, "「网站 24 小时在线，演讲后欢迎扫码或点链接体验。GitHub 仓库源代码、提交历史也都开源在右上角链接里。」")
script(doc, "「Thank You！」")

# ============================================================
# Q&A Prep
# ============================================================
doc.add_page_break()
heading1(doc, "Q&A 准备 — 3 个高概率问题 + 标准答案")

heading2(doc, "Q1：你们的相关性公式 40/30/30 是怎么定的？")
body(doc, "回答（陈瀚中）：这是迭代调出来的。最早版本是 50/30/20，文本匹配权重最高。但跑下来发现：标题完全匹配的冷门资源容易抢前列，而真正高质量的资源因为标题用了同义词反而排不上去。")
body(doc, "改成 40/30/30 后，把质量信号（下载热度 + 评分）的总权重提到 60%，让用户行为投票出来的优质资源浮上来。如果未来真上线，可以做 A/B 测试持续优化这三个权重。")

heading2(doc, "Q2：积分系统会不会被刷？比如自己注册多个账号互相下载？")
body(doc, "回答（郁凯杰）：好问题。我们做了三层防御：")
bullet(doc, "数据库层 UNIQUE 约束：一个用户对一个资源只能评 1 次")
bullet(doc, "积分引擎规则：自评不奖励、自己下载不扣分、4 星以下不发 +1、每日免费下载上限")
bullet(doc, "Admin 审核：所有上传必须经过审核才能进入流通，低质或滥用会被拒绝")
body(doc, "更深层的多人共谋我们没做，因为校园场景账号要学号实名，违规成本很高。如果上线发现问题，可以再加风控规则。")

heading2(doc, "Q3：你们的部署是真的吗？还是 PPT 截图？")
body(doc, "回答（连宇翔 / 郁凯杰）：是真的，这就是我们一开始当场注册账号的原因。")
bullet(doc, "数据库：MariaDB 真在云服务器上跑，注册的账号已经写入 users 表，演示完可远程登入服务器查看")
bullet(doc, "API：FastAPI 在 Docker 容器里 24/7 响应，所有演示动作都打到真实后端")
bullet(doc, "前端：Prototype.html 是真用户能访问的网站，不是 PPT 截图")
bullet(doc, "压测：100 并发请求 0 错误，1000 笔积分扣减 0 race condition，全是真实测出来的，不是设计预测")
body(doc, "演讲结束后平台还会继续在线，老师同学可以随时回访体验。")

heading2(doc, "Q4（备用）：为什么选 MariaDB 不是 MySQL 8 或 PostgreSQL？")
body(doc, "回答（郁凯杰）：MariaDB 10.11 是 MySQL 的开源分叉，SQL 完全兼容，但内存占用更紧凑（180MB 上限完全够用），也避免 Oracle 系商业许可问题。Postgres 在 1GB VPS 上也可以跑，但我们的查询模式偏 OLTP 简单事务，MariaDB 已足够，没必要为可能用不到的高级特性多花资源。")

# === Tips ===
doc.add_page_break()
heading1(doc, "舞台呈现 Tips")

heading2(doc, "时间控制（最重要）")
bullet(doc, "Demo 6 分钟是核心——切忌时间不够时跳过 Demo，宁可砍 reflection 那一页")
bullet(doc, "如果 Demo 顺利只用了 4 分钟，多出 2 分钟给老师互动注册、提问")
bullet(doc, "用手表 / 手机定时，每张 slide 大概多少时间心里有数")

heading2(doc, "演示路径绝对不能错")
bullet(doc, "演讲开始前 5 分钟，浏览器里测一遍：登录 → 搜索 → 下载 → 流水")
bullet(doc, "URL 提前发到自己手机备忘 + 老师电脑（万一笔记本挂了能切换）")
bullet(doc, "Demo 时**字号调到 1.3x**，让后排能看清")

heading2(doc, "应急预案")
bullet(doc, "Cloudflare URL 失效（quick tunnel 偶尔会变）→ 备用 IP：http://178.157.59.239/")
bullet(doc, "Demo 完全卡死 → 用 PPT 上的截图过一遍流程，说明真实行为")
bullet(doc, "投影仪故障 → PPT 转成 PDF 提前存 USB + iCloud + 老师邮箱三处")
bullet(doc, "时间不够 → 跳过 Slide 9 反思，直接从 Slide 8 跳到 Slide 10 Q&A")

heading2(doc, "讲话节奏")
bullet(doc, "数字要慢念两遍：『87 毫秒，比目标快 22 倍』")
bullet(doc, "切换讲者时明确说『把话筒交给 X』+ 微笑握手 1 秒")
bullet(doc, "如果忘词，停 1 秒看 PPT 再继续，比说『嗯额啊』强一万倍")

doc.add_paragraph()
footer = doc.add_paragraph("End of Presentation Script v2  |  15-min lean version  |  May 2026")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save('/Users/yuxianglian/Documents/系统分析与设计/SAD_Project/Presentation_Script.docx')
print("Done: Presentation_Script.docx (15-min version)")
