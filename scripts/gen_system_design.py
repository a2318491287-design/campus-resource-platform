from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

def heading1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(15)
    return p

def code_block(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(9)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3 + level*0.25)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F497D')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'D6E4F0')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ===================== COVER PAGE =====================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("SYSTEM DESIGN DOCUMENT")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Campus Academic Resource Sharing Platform")
r.font.size = Pt(16)
r.font.bold = True

doc.add_paragraph()
info = [
    ("Project Title:", "Campus Academic Resource Sharing Platform Development"),
    ("Document Version:", "1.0"),
    ("Prepared by:", "Lian Yuxiang (1230020693) / Yu Kaijie (1230020426)"),
    ("Team Members:", "Lian Yuxiang 1230020693  |  Yu Kaijie 1230020426  |  Chen Hanzhong 1230032209"),
    ("Course:", "System Analysis and Design"),
    ("Lecturer:", "Dr. CHE Pak Hou (Howard)"),
    ("Date:", "May 2, 2026"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}  ")
    r.font.size = Pt(11)
    r.font.bold = True
    r2 = p.add_run(value)
    r2.font.size = Pt(11)

doc.add_page_break()

# ===================== TOC =====================
heading1(doc, "Table of Contents")
toc = [
    "1. Introduction ................................................... 3",
    "2. System Architecture Design ..................................... 3",
    "3. Module Design .................................................. 5",
    "4. Class Diagram .................................................. 6",
    "5. Sequence Diagrams .............................................. 8",
    "6. State Diagram .................................................. 11",
    "7. Component Diagram .............................................. 12",
    "8. Database Design (MySQL) ........................................ 13",
    "9. Interface Design Specifications ................................ 16",
    "10. Security Design ............................................... 17",
    "11. Design Decision Log ........................................... 18",
]
for t in toc:
    p = doc.add_paragraph(t)
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
doc.add_page_break()

# ===================== SECTION 1 =====================
heading1(doc, "1. Introduction")
heading2(doc, "1.1 Purpose")
body(doc, "This System Design Document (SDD) translates the requirements defined in the Requirements Analysis Document (RAD v1.0) into a concrete architectural and component-level design for the Campus Academic Resource Sharing Platform. It serves as the authoritative guide for prototype development, database implementation, and testing, following ISO 12207 software lifecycle standards.")

heading2(doc, "1.2 Design Principles")
bullet(doc, "Separation of Concerns: Presentation, business logic, and data layers are strictly separated.")
bullet(doc, "Single Responsibility: Each module handles one clearly defined responsibility.")
bullet(doc, "Simplicity Over Cleverness: Straightforward solutions preferred to minimize maintenance burden on a 3-person team.")
bullet(doc, "Security by Design: Authentication, authorization, and data validation are built into the architecture, not retrofitted.")
bullet(doc, "Traceability: Every design component traces back to one or more functional requirements (FR-XX).")

heading2(doc, "1.3 Referenced Documents")
bullet(doc, "Requirements Analysis Document v1.0 — April 4, 2026")
bullet(doc, "Project Charter v1.0 — March 15, 2026")
bullet(doc, "Tilley & Rosenblatt, Systems Analysis and Design, 12th Edition")
bullet(doc, "ISO/IEC 12207:2017 — Software Lifecycle Processes")

# ===================== SECTION 2 =====================
doc.add_paragraph()
heading1(doc, "2. System Architecture Design")
heading2(doc, "2.1 Architecture Style")
body(doc, "The platform adopts a Three-Tier Architecture pattern, clearly separating the Presentation Tier, Business Logic Tier, and Data Tier. This approach supports independent development of each layer and aligns with the team's available tools (Figma for presentation, Python/Node.js for logic, MySQL for data).")

heading2(doc, "2.2 Three-Tier Architecture Overview")
body(doc, "Tier 1 — Presentation Layer (Client Side):")
bullet(doc, "Technology: HTML5 / CSS3 / JavaScript (React or Vue.js — TBD at implementation phase)")
bullet(doc, "Responsibility: Render UI, capture user inputs, display search results, handle client-side validation")
bullet(doc, "Tools: Figma (prototype), Chrome DevTools (testing)")

body(doc, "Tier 2 — Business Logic Layer (Server Side):")
bullet(doc, "Technology: Python (Flask/FastAPI) or Node.js (Express) — TBD")
bullet(doc, "Responsibility: Authentication, search engine, points calculation, file management, notification dispatch")
bullet(doc, "Key Modules: AuthModule, SearchModule, PointsEngine, UploadManager, AdminModule")

body(doc, "Tier 3 — Data Layer:")
bullet(doc, "Technology: MySQL 8.0")
bullet(doc, "Responsibility: Persistent storage of users, resources, tags, ratings, point records")
bullet(doc, "File Storage: Local file system (prototype) / Cloud object storage (production, out of scope)")

heading2(doc, "2.3 Architecture Diagram (Textual Representation)")
body(doc, "[ To be rendered in draw.io — description below ]", indent=True)
for line in [
    "┌─────────────────────────────────────────────────────────┐",
    "│                  PRESENTATION TIER                       │",
    "│   Web Browser / Mobile Browser                          │",
    "│   Pages: Search | Upload | Profile | Points | Admin     │",
    "└──────────────────────┬──────────────────────────────────┘",
    "                       │  HTTP / HTTPS (REST API calls)   ",
    "┌──────────────────────▼──────────────────────────────────┐",
    "│                 BUSINESS LOGIC TIER                      │",
    "│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌───────────┐  │",
    "│  │   Auth   │ │  Search  │ │  Points  │ │  Upload   │  │",
    "│  │  Module  │ │  Module  │ │  Engine  │ │  Manager  │  │",
    "│  └──────────┘ └──────────┘ └──────────┘ └───────────┘  │",
    "│  ┌──────────────────────────────────────────────────┐   │",
    "│  │          Admin Module  |  Notification Service   │   │",
    "│  └──────────────────────────────────────────────────┘   │",
    "└──────────────────────┬──────────────────────────────────┘",
    "                       │  SQL Queries (MySQL Connector)   ",
    "┌──────────────────────▼──────────────────────────────────┐",
    "│                    DATA TIER                             │",
    "│  MySQL 8.0 Database                                     │",
    "│  Tables: users | resources | tags | resource_tags |     │",
    "│          ratings | point_records | downloads |          │",
    "│          redemptions                                    │",
    "└─────────────────────────────────────────────────────────┘",
]:
    code_block(doc, line)

# ===================== SECTION 3 =====================
doc.add_paragraph()
heading1(doc, "3. Module Design")
mod_headers = ["Module Name", "Responsibility", "Key Operations", "Depends On", "Related FRs"]
mod_rows = [
    ["AuthModule", "User registration, login, session management, password reset", "register(), login(), logout(), resetPassword(), validateSession()", "UserDAO (Data Layer)", "FR-01, FR-02, FR-04"],
    ["SearchModule", "Execute multi-dimensional resource search and relevance ranking", "search(keyword, filters), computeRelevanceScore(), getPreviewSnippet(), getSearchHistory()", "ResourceDAO, TagDAO", "FR-05, FR-06, FR-07, FR-08, FR-09, FR-10"],
    ["UploadManager", "Handle file uploads, validate format/size, trigger review queue entry", "uploadFile(), validateMetadata(), queueForReview(), updateStatus()", "ResourceDAO, TagDAO, PointsEngine", "FR-20, FR-21, FR-22"],
    ["PointsEngine", "Calculate, award, deduct, and record all point transactions", "awardPoints(userId, amount, reason), deductPoints(), getBalance(), getHistory()", "PointRecordDAO, UserDAO", "FR-11 to FR-19"],
    ["RatingModule", "Accept and store user ratings and comments, update resource avg_rating", "submitRating(), getResourceRatings(), computeAverageRating()", "RatingDAO, PointsEngine", "FR-23"],
    ["AdminModule", "Manage review queue, approve/reject uploads, handle reports", "listPendingResources(), approveResource(), rejectResource(), handleReport()", "ResourceDAO, PointsEngine, NotifService", "FR-22, FR-25"],
    ["RedemptionModule", "Process point redemption requests, activate rewards", "listRedeemOptions(), redeem(userId, optionId), activateReward()", "PointsEngine, UserDAO", "FR-15, FR-16"],
    ["NotifService", "Send in-app notifications for key events", "notify(userId, message, type)", "UserDAO", "FR-19"],
    ["LeaderboardModule", "Compute and cache monthly top contributors ranking", "getMonthlyLeaderboard(), computeContributorScore()", "PointRecordDAO", "FR-17"],
]
add_table(doc, mod_headers, mod_rows, [1.4, 1.5, 1.8, 1.2, 1.1])

# ===================== SECTION 4: CLASS DIAGRAM =====================
doc.add_paragraph()
heading1(doc, "4. Class Diagram")
heading2(doc, "4.1 Overview")
body(doc, "The class diagram below describes the core domain classes, their attributes, methods, and relationships. It directly maps to the MySQL table design in Section 8.")

heading2(doc, "4.2 Core Classes")

classes = [
    ("User",
     ["- userId: int", "- studentId: String", "- username: String", "- passwordHash: String", "- email: String", "- pointsBalance: int", "- uploadCount: int", "- downloadCredits: int", "- isAdmin: boolean", "- createdAt: DateTime"],
     ["+ register(): boolean", "+ login(studentId, password): String", "+ logout(): void", "+ getProfile(): UserProfile", "+ getPointsBalance(): int", "+ updateDownloadCredits(amount: int): void"]
    ),
    ("Resource",
     ["- resourceId: int", "- title: String", "- description: String", "- filePath: String", "- fileType: String (PDF/DOCX/PPTX/IMAGE)", "- fileSize: long", "- courseCode: String", "- academicYear: int", "- resourceType: String", "- status: ResourceStatus", "- avgRating: float", "- downloadCount: int", "- uploaderId: int", "- createdAt: DateTime"],
     ["+ getDetail(): ResourceDetail", "+ updateMetadata(title, desc, tags): void", "+ incrementDownloadCount(): void", "+ updateAvgRating(): void", "+ changeStatus(status: ResourceStatus): void"]
    ),
    ("Tag",
     ["- tagId: int", "- tagName: String", "- category: TagCategory (COURSE/TYPE/KEYWORD)"],
     ["+ findByName(name: String): Tag", "+ findByCategory(cat: TagCategory): List<Tag>"]
    ),
    ("Rating",
     ["- ratingId: int", "- resourceId: int", "- userId: int", "- stars: int", "- comment: String", "- createdAt: DateTime"],
     ["+ submit(resourceId, userId, stars, comment): Rating", "+ getByResource(resourceId): List<Rating>"]
    ),
    ("PointRecord",
     ["- recordId: int", "- userId: int", "- resourceId: int (nullable)", "- actionType: PointActionType", "- pointsDelta: int", "- balanceAfter: int", "- createdAt: DateTime"],
     ["+ record(userId, delta, action, resourceId): PointRecord", "+ getHistoryByUser(userId): List<PointRecord>"]
    ),
    ("PointsEngine",
     ["- «static» UPLOAD_REWARD: int = 10", "- «static» DOWNLOAD_RECEIVED_REWARD: int = 2", "- «static» RATING_RECEIVED_REWARD: int = 1", "- «static» DOWNLOAD_COST: int = 5"],
     ["+ awardUpload(userId): void", "+ awardDownloadReceived(uploaderId): void", "+ awardRatingReceived(uploaderId): void", "+ chargeDownload(userId): boolean", "+ redeem(userId, option: RedeemOption): boolean"]
    ),
    ("SearchEngine",
     ["- «static» WEIGHT_MATCH: float = 0.40", "- «static» WEIGHT_DOWNLOADS: float = 0.30", "- «static» WEIGHT_RATING: float = 0.30"],
     ["+ search(keyword: String, filters: SearchFilter): List<ResourceResult>", "+ computeRelevanceScore(resource, keyword): float", "+ getPreview(resourceId): String", "+ getRecommendations(resourceId): List<Resource>"]
    ),
    ("SearchFilter",
     ["- courseCode: String (optional)", "- academicYear: int (optional)", "- resourceType: String (optional)", "- minRating: float (optional)"],
     ["+ isValid(): boolean"]
    ),
]

for cls_name, attrs, methods in classes:
    heading3(doc, f"Class: {cls_name}")
    table = doc.add_table(rows=3, cols=1)
    table.style = 'Table Grid'
    # Class name row
    name_cell = table.rows[0].cells[0]
    name_cell.text = cls_name
    name_cell.paragraphs[0].runs[0].font.bold = True
    name_cell.paragraphs[0].runs[0].font.size = Pt(11)
    name_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = name_cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '2E74B5')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    name_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    # Attributes row
    attr_cell = table.rows[1].cells[0]
    attr_cell.text = "\n".join(attrs)
    attr_cell.paragraphs[0].runs[0].font.size = Pt(9)
    attr_cell.paragraphs[0].runs[0].font.name = 'Courier New'
    # Methods row
    meth_cell = table.rows[2].cells[0]
    meth_cell.text = "\n".join(methods)
    meth_cell.paragraphs[0].runs[0].font.size = Pt(9)
    meth_cell.paragraphs[0].runs[0].font.name = 'Courier New'
    doc.add_paragraph()

heading2(doc, "4.3 Enumerations")
enum_data = [
    ("ResourceStatus", ["PENDING", "PUBLISHED", "REJECTED", "REMOVED"]),
    ("PointActionType", ["UPLOAD_APPROVED", "DOWNLOAD_RECEIVED", "RATING_RECEIVED", "SPEND_DOWNLOAD", "REDEEM_DOWNLOAD_CREDIT", "REDEEM_PIN"]),
    ("TagCategory", ["COURSE", "TYPE", "KEYWORD"]),
    ("RedeemOption", ["DOWNLOAD_CREDIT_10 (cost: 50 pts)", "RESOURCE_PIN_7DAYS (cost: 100 pts)"]),
]
for enum_name, values in enum_data:
    p = doc.add_paragraph()
    r = p.add_run(f"«enumeration» {enum_name}: ")
    r.font.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(" | ".join(values))
    r2.font.size = Pt(11)

heading2(doc, "4.4 Class Relationships Summary")
rel_headers = ["From Class", "Relationship", "To Class", "Multiplicity", "Notes"]
rel_rows = [
    ["User", "uploads →", "Resource", "1 : 0..*", "uploaderId FK"],
    ["User", "submits →", "Rating", "1 : 0..*", "one rating per user per resource (enforced by UNIQUE constraint)"],
    ["User", "has →", "PointRecord", "1 : 0..*", "all earning/spending events"],
    ["Resource", "described by →", "Tag", "M : M", "via ResourceTag junction"],
    ["Resource", "has →", "Rating", "1 : 0..*", ""],
    ["PointsEngine", "uses →", "PointRecord", "creates", "factory pattern"],
    ["PointsEngine", "modifies →", "User", "pointsBalance", "atomic update"],
    ["SearchEngine", "queries →", "Resource", "reads", "no direct ownership"],
    ["SearchEngine", "uses →", "SearchFilter", "parameter", "value object"],
]
add_table(doc, rel_headers, rel_rows, [1.4, 1.1, 1.1, 1.0, 2.4])

# ===================== SECTION 5: SEQUENCE DIAGRAMS =====================
doc.add_paragraph()
heading1(doc, "5. Sequence Diagrams")
body(doc, "Sequence diagrams describe the time-ordered interactions between objects for the three most critical use cases. These should be rendered in draw.io or PlantUML using the textual descriptions below.")

heading2(doc, "5.1 Sequence Diagram: Search and Download Resource (UC-03 + UC-04)")
body(doc, "Actors/Objects: Student, SearchPage (UI), SearchEngine, ResourceDAO, PointsEngine, PointRecordDAO, FileStorage")
body(doc, "PlantUML notation:")
seq1 = """@startuml
actor Student
participant "SearchPage\\n(UI)" as UI
participant SearchEngine
participant ResourceDAO
participant PointsEngine
participant FileStorage

Student -> UI : enter keyword + filters
UI -> SearchEngine : search(keyword, filters)
SearchEngine -> ResourceDAO : queryResources(keyword, filters)
ResourceDAO --> SearchEngine : List<Resource>
SearchEngine -> SearchEngine : computeRelevanceScore(resources)
SearchEngine --> UI : ranked List<ResourceResult>
UI --> Student : display search results

Student -> UI : click Download on result
UI -> PointsEngine : checkBalance(userId)
PointsEngine --> UI : balance = 35

alt balance >= 5 OR free downloads remain
  UI -> PointsEngine : chargeDownload(userId)
  PointsEngine -> ResourceDAO : incrementDownloadCount(resourceId)
  PointsEngine -> PointRecordDAO : record(userId, -5, SPEND_DOWNLOAD)
  PointsEngine --> UI : success
  UI -> FileStorage : getFile(resourceId)
  FileStorage --> UI : fileStream
  UI --> Student : file downloaded
else insufficient points
  UI --> Student : show "Insufficient Points" dialog
end

PointsEngine -> ResourceDAO : getUploaderId(resourceId)
ResourceDAO --> PointsEngine : uploaderId
PointsEngine -> PointsEngine : awardDownloadReceived(uploaderId)
PointsEngine -> PointRecordDAO : record(uploaderId, +2, DOWNLOAD_RECEIVED)
@enduml"""
for line in seq1.strip().split('\n'):
    code_block(doc, line)

heading2(doc, "5.2 Sequence Diagram: Upload Resource (UC-05)")
body(doc, "Actors/Objects: Student, UploadPage (UI), UploadManager, ResourceDAO, TagDAO, AdminModule, PointsEngine, NotifService")
seq2 = """@startuml
actor Student
actor Admin
participant "UploadPage\\n(UI)" as UI
participant UploadManager
participant ResourceDAO
participant TagDAO
participant AdminModule
participant PointsEngine
participant NotifService

Student -> UI : select file + fill metadata form
UI -> UploadManager : uploadFile(file, metadata)
UploadManager -> UploadManager : validateFileFormat(file)
UploadManager -> UploadManager : validateFileSize(file)

alt valid file
  UploadManager -> ResourceDAO : createResource(metadata, status=PENDING)
  ResourceDAO --> UploadManager : resourceId
  UploadManager -> TagDAO : linkTags(resourceId, tags)
  UploadManager -> AdminModule : addToReviewQueue(resourceId)
  UploadManager --> UI : "Submitted for review"
  UI --> Student : show confirmation message
else invalid file
  UploadManager --> UI : error (format/size)
  UI --> Student : show error message
end

Admin -> AdminModule : reviewResource(resourceId, decision=APPROVE)
AdminModule -> ResourceDAO : updateStatus(resourceId, PUBLISHED)
AdminModule -> PointsEngine : awardUpload(uploaderId)
PointsEngine -> ResourceDAO : getUploaderId(resourceId)
PointsEngine -> PointsEngine : creditPoints(uploaderId, +10)
PointsEngine --> AdminModule : success
AdminModule -> NotifService : notify(uploaderId, "Resource approved! +10 pts")
NotifService --> Student : in-app notification
@enduml"""
for line in seq2.strip().split('\n'):
    code_block(doc, line)

heading2(doc, "5.3 Sequence Diagram: Redeem Points (UC-08)")
body(doc, "Actors/Objects: Student, PointsDashboard (UI), RedemptionModule, PointsEngine, PointRecordDAO")
seq3 = """@startuml
actor Student
participant "PointsDashboard\\n(UI)" as UI
participant RedemptionModule
participant PointsEngine
participant PointRecordDAO

Student -> UI : navigate to Points Dashboard
UI -> PointsEngine : getBalance(userId)
PointsEngine --> UI : balance = 75
UI --> Student : show balance + redemption options

Student -> UI : select "50 pts → 10 downloads"
UI --> Student : show confirmation dialog (cost: 50 pts, benefit: 10 downloads)
Student -> UI : confirm redemption

UI -> RedemptionModule : redeem(userId, DOWNLOAD_CREDIT_10)
RedemptionModule -> PointsEngine : checkBalance(userId)
PointsEngine --> RedemptionModule : balance = 75

alt balance >= 50
  RedemptionModule -> PointsEngine : deductPoints(userId, 50, REDEEM_DOWNLOAD_CREDIT)
  PointsEngine -> PointRecordDAO : record(userId, -50, REDEEM_DOWNLOAD_CREDIT)
  PointsEngine -> PointsEngine : updateDownloadCredits(userId, +10)
  PointsEngine --> RedemptionModule : success (newBalance=25, newCredits=10)
  RedemptionModule --> UI : success
  UI --> Student : "Redemption successful! Balance: 25 pts, Download Credits: 10"
else balance < 50 (race condition guard)
  RedemptionModule --> UI : insufficient balance
  UI --> Student : show error
end
@enduml"""
for line in seq3.strip().split('\n'):
    code_block(doc, line)

# ===================== SECTION 6: STATE DIAGRAM =====================
doc.add_paragraph()
heading1(doc, "6. State Diagram")
heading2(doc, "6.1 Resource Lifecycle State Diagram")
body(doc, "A Resource object transitions through the following states from creation to removal:")
state_lines = [
    "@startuml",
    "[*] --> PENDING : User submits upload",
    "PENDING --> PUBLISHED : Admin approves",
    "PENDING --> REJECTED : Admin rejects",
    "PUBLISHED --> REMOVED : Admin removes (report upheld)",
    "PUBLISHED --> PUBLISHED : User edits metadata",
    "REJECTED --> PENDING : User resubmits (after corrections)",
    "REMOVED --> [*]",
    "@enduml",
]
for line in state_lines:
    code_block(doc, line)

body(doc, "\nState Descriptions:")
sd_headers = ["State", "Description", "Entry Action", "Exit Action"]
sd_rows = [
    ["PENDING", "Resource uploaded but awaiting admin review. Visible only to uploader in 'My Uploads'.", "Send notification to uploader: 'Under review'", "Admin makes approval decision"],
    ["PUBLISHED", "Resource approved and publicly searchable. Contributes to uploader's points when downloaded.", "Award 10 pts to uploader; send notification; add to search index", "Admin removes or user deletes"],
    ["REJECTED", "Resource failed review (copyright issue, incorrect metadata, irrelevant content).", "Send rejection notification with reason; no points awarded", "User corrects and resubmits"],
    ["REMOVED", "Resource removed due to policy violation after being published. Permanently unavailable.", "Send removal notification; deduct 5 pts as penalty", "Terminal state — archived in DB"],
]
add_table(doc, sd_headers, sd_rows, [1.1, 2.3, 1.8, 1.8])

heading2(doc, "6.2 Points Transaction State (User Balance)")
body(doc, "The user's points balance is not a state machine but follows these invariants:")
bullet(doc, "Balance can never go below 0. If a SPEND action would result in negative balance, the transaction is blocked (except for the 3 free daily downloads policy).")
bullet(doc, "All balance changes are atomic: PointRecord is written and balance updated in the same database transaction (BEGIN/COMMIT).")
bullet(doc, "Balance is denormalized on the User record for fast reads; PointRecord table is the authoritative source of truth for audits.")

# ===================== SECTION 7: COMPONENT DIAGRAM =====================
doc.add_paragraph()
heading1(doc, "7. Component Diagram")
body(doc, "The component diagram shows the physical packaging of the system into deployable units and their dependencies:")
comp_lines = [
    "@startuml",
    "package \"Client Browser\" {",
    "  [SearchPage]",
    "  [UploadPage]",
    "  [ProfilePage]",
    "  [PointsDashboard]",
    "  [AdminPanel]",
    "}",
    "",
    "package \"Application Server\" {",
    "  [AuthModule]",
    "  [SearchModule]",
    "  [UploadManager]",
    "  [PointsEngine]",
    "  [RedemptionModule]",
    "  [AdminModule]",
    "  [NotifService]",
    "  [LeaderboardModule]",
    "  [RESTful API Gateway]",
    "}",
    "",
    "package \"Data Layer\" {",
    "  database \"MySQL 8.0\" {",
    "    [users]",
    "    [resources]",
    "    [tags / resource_tags]",
    "    [ratings]",
    "    [point_records]",
    "    [downloads]",
    "    [redemptions]",
    "  }",
    "  [File Storage]",
    "}",
    "",
    "[Client Browser] --> [RESTful API Gateway] : HTTPS",
    "[RESTful API Gateway] --> [AuthModule]",
    "[RESTful API Gateway] --> [SearchModule]",
    "[RESTful API Gateway] --> [UploadManager]",
    "[RESTful API Gateway] --> [PointsEngine]",
    "[UploadManager] --> [File Storage]",
    "[SearchModule] --> [MySQL 8.0]",
    "[PointsEngine] --> [MySQL 8.0]",
    "[AuthModule] --> [MySQL 8.0]",
    "@enduml",
]
for line in comp_lines:
    code_block(doc, line)

# ===================== SECTION 8: DATABASE DESIGN =====================
doc.add_paragraph()
heading1(doc, "8. Database Design (MySQL 8.0)")
heading2(doc, "8.1 Entity Overview")
body(doc, "The following 8 tables implement the ERD defined in the Requirements Analysis Document:")

db_headers = ["Table Name", "Primary Key", "Description", "Rows Est."]
db_rows = [
    ["users", "user_id (INT AUTO_INCREMENT)", "Student user accounts, credentials, points balance", "~5,000"],
    ["resources", "resource_id (INT AUTO_INCREMENT)", "All academic resource files and metadata", "~50,000"],
    ["tags", "tag_id (INT AUTO_INCREMENT)", "Searchable tag vocabulary (course codes, types, keywords)", "~500"],
    ["resource_tags", "(resource_id, tag_id) composite", "Many-to-many linking table for resources and tags", "~200,000"],
    ["ratings", "rating_id (INT AUTO_INCREMENT)", "User ratings and comments for resources", "~30,000"],
    ["point_records", "record_id (INT AUTO_INCREMENT)", "Complete audit log of all point transactions", "~500,000"],
    ["downloads", "download_id (INT AUTO_INCREMENT)", "Log of every download event for analytics", "~100,000"],
    ["redemptions", "redemption_id (INT AUTO_INCREMENT)", "Records of all point redemption activations", "~10,000"],
]
add_table(doc, db_headers, db_rows, [1.5, 2.1, 2.5, 0.9])

heading2(doc, "8.2 Complete SQL Schema")
sql_schema = """-- ================================================================
-- Campus Academic Resource Sharing Platform
-- MySQL 8.0 Database Schema  |  Version 1.0  |  May 2, 2026
-- ================================================================

CREATE DATABASE IF NOT EXISTS campus_resource_platform
  DEFAULT CHARACTER SET utf8mb4
  DEFAULT COLLATE utf8mb4_unicode_ci;

USE campus_resource_platform;

-- ----------------------------------------------------------------
-- Table: users
-- ----------------------------------------------------------------
CREATE TABLE users (
  user_id          INT          NOT NULL AUTO_INCREMENT,
  student_id       VARCHAR(20)  NOT NULL UNIQUE,
  username         VARCHAR(50)  NOT NULL,
  password_hash    VARCHAR(255) NOT NULL,          -- bcrypt, cost=12
  email            VARCHAR(100) NOT NULL UNIQUE,
  points_balance   INT          NOT NULL DEFAULT 0 CHECK (points_balance >= 0),
  upload_count     INT          NOT NULL DEFAULT 0,
  download_credits INT          NOT NULL DEFAULT 3, -- 3 free daily resets at midnight
  is_admin         TINYINT(1)   NOT NULL DEFAULT 0,
  created_at       DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP,
  updated_at       DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP
                               ON UPDATE CURRENT_TIMESTAMP,
  PRIMARY KEY (user_id),
  INDEX idx_student_id (student_id),
  INDEX idx_email (email)
) ENGINE=InnoDB;

-- ----------------------------------------------------------------
-- Table: resources
-- ----------------------------------------------------------------
CREATE TABLE resources (
  resource_id    INT          NOT NULL AUTO_INCREMENT,
  title          VARCHAR(200) NOT NULL,
  description    TEXT,
  file_path      VARCHAR(500) NOT NULL,
  file_type      ENUM('PDF','DOCX','PPTX','IMAGE') NOT NULL,
  file_size      BIGINT       NOT NULL,            -- bytes
  course_code    VARCHAR(20)  NOT NULL,
  academic_year  YEAR         NOT NULL,
  resource_type  ENUM('NOTES','PAST_PAPER','ASSIGNMENT','OTHER') NOT NULL,
  status         ENUM('PENDING','PUBLISHED','REJECTED','REMOVED')
                              NOT NULL DEFAULT 'PENDING',
  avg_rating     DECIMAL(3,2) DEFAULT NULL,
  download_count INT          NOT NULL DEFAULT 0,
  uploader_id    INT          NOT NULL,
  rejection_reason VARCHAR(500) DEFAULT NULL,
  created_at     DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP,
  updated_at     DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP
                              ON UPDATE CURRENT_TIMESTAMP,
  PRIMARY KEY (resource_id),
  FOREIGN KEY (uploader_id) REFERENCES users(user_id) ON DELETE CASCADE,
  FULLTEXT INDEX ft_search (title, description),
  INDEX idx_course_code (course_code),
  INDEX idx_status (status),
  INDEX idx_academic_year (academic_year)
) ENGINE=InnoDB;

-- ----------------------------------------------------------------
-- Table: tags
-- ----------------------------------------------------------------
CREATE TABLE tags (
  tag_id    INT          NOT NULL AUTO_INCREMENT,
  tag_name  VARCHAR(50)  NOT NULL UNIQUE,
  category  ENUM('COURSE','TYPE','KEYWORD') NOT NULL,
  PRIMARY KEY (tag_id),
  INDEX idx_category (category)
) ENGINE=InnoDB;

-- ----------------------------------------------------------------
-- Table: resource_tags  (junction table)
-- ----------------------------------------------------------------
CREATE TABLE resource_tags (
  resource_id INT NOT NULL,
  tag_id      INT NOT NULL,
  PRIMARY KEY (resource_id, tag_id),
  FOREIGN KEY (resource_id) REFERENCES resources(resource_id) ON DELETE CASCADE,
  FOREIGN KEY (tag_id)      REFERENCES tags(tag_id)           ON DELETE CASCADE
) ENGINE=InnoDB;

-- ----------------------------------------------------------------
-- Table: ratings
-- ----------------------------------------------------------------
CREATE TABLE ratings (
  rating_id   INT       NOT NULL AUTO_INCREMENT,
  resource_id INT       NOT NULL,
  user_id     INT       NOT NULL,
  stars       TINYINT   NOT NULL CHECK (stars BETWEEN 1 AND 5),
  comment     TEXT,
  created_at  DATETIME  NOT NULL DEFAULT CURRENT_TIMESTAMP,
  PRIMARY KEY (rating_id),
  UNIQUE KEY uq_user_resource (user_id, resource_id),  -- one rating per user
  FOREIGN KEY (resource_id) REFERENCES resources(resource_id) ON DELETE CASCADE,
  FOREIGN KEY (user_id)     REFERENCES users(user_id)         ON DELETE CASCADE
) ENGINE=InnoDB;

-- ----------------------------------------------------------------
-- Table: point_records
-- ----------------------------------------------------------------
CREATE TABLE point_records (
  record_id     INT     NOT NULL AUTO_INCREMENT,
  user_id       INT     NOT NULL,
  resource_id   INT     DEFAULT NULL,
  action_type   ENUM('UPLOAD_APPROVED','DOWNLOAD_RECEIVED',
                     'RATING_RECEIVED','SPEND_DOWNLOAD',
                     'REDEEM_DOWNLOAD_CREDIT','REDEEM_PIN') NOT NULL,
  points_delta  INT     NOT NULL,  -- positive = earn, negative = spend
  balance_after INT     NOT NULL,
  created_at    DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
  PRIMARY KEY (record_id),
  FOREIGN KEY (user_id)     REFERENCES users(user_id)         ON DELETE CASCADE,
  FOREIGN KEY (resource_id) REFERENCES resources(resource_id) ON DELETE SET NULL,
  INDEX idx_user_id (user_id),
  INDEX idx_created_at (created_at)
) ENGINE=InnoDB;

-- ----------------------------------------------------------------
-- Table: downloads
-- ----------------------------------------------------------------
CREATE TABLE downloads (
  download_id  INT      NOT NULL AUTO_INCREMENT,
  resource_id  INT      NOT NULL,
  user_id      INT      NOT NULL,
  downloaded_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
  PRIMARY KEY (download_id),
  FOREIGN KEY (resource_id) REFERENCES resources(resource_id) ON DELETE CASCADE,
  FOREIGN KEY (user_id)     REFERENCES users(user_id)         ON DELETE CASCADE,
  INDEX idx_resource_id (resource_id),
  INDEX idx_user_id (user_id)
) ENGINE=InnoDB;

-- ----------------------------------------------------------------
-- Table: redemptions
-- ----------------------------------------------------------------
CREATE TABLE redemptions (
  redemption_id INT      NOT NULL AUTO_INCREMENT,
  user_id       INT      NOT NULL,
  reward_type   ENUM('DOWNLOAD_CREDIT','PIN') NOT NULL,
  points_cost   INT      NOT NULL,
  resource_id   INT      DEFAULT NULL,  -- required for PIN reward
  activated_at  DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
  expires_at    DATETIME DEFAULT NULL,
  PRIMARY KEY (redemption_id),
  FOREIGN KEY (user_id)     REFERENCES users(user_id)         ON DELETE CASCADE,
  FOREIGN KEY (resource_id) REFERENCES resources(resource_id) ON DELETE SET NULL
) ENGINE=InnoDB;"""

for line in sql_schema.strip().split('\n'):
    code_block(doc, line)

heading2(doc, "8.3 Key Queries")
heading3(doc, "Q1: Relevance-Ranked Search")
q1 = """SELECT r.resource_id, r.title, r.course_code, r.resource_type,
       r.academic_year, r.avg_rating, r.download_count,
       (
         MATCH(r.title, r.description) AGAINST (:keyword IN NATURAL LANGUAGE MODE) * 0.40
         + (r.download_count / 1000.0) * 0.30
         + (COALESCE(r.avg_rating, 0) / 5.0) * 0.30
       ) AS relevance_score
FROM resources r
WHERE r.status = 'PUBLISHED'
  AND (:courseCode IS NULL OR r.course_code = :courseCode)
  AND (:year IS NULL OR r.academic_year = :year)
  AND (:type IS NULL OR r.resource_type = :type)
  AND (:minRating IS NULL OR r.avg_rating >= :minRating)
ORDER BY relevance_score DESC
LIMIT 20 OFFSET :offset;"""
for line in q1.strip().split('\n'):
    code_block(doc, line)

heading3(doc, "Q2: Monthly Leaderboard")
q2 = """SELECT u.username, SUM(pr.points_delta) AS monthly_earned
FROM point_records pr
JOIN users u ON pr.user_id = u.user_id
WHERE pr.action_type IN ('UPLOAD_APPROVED','DOWNLOAD_RECEIVED','RATING_RECEIVED')
  AND pr.created_at >= DATE_FORMAT(NOW(), '%Y-%m-01')
GROUP BY pr.user_id, u.username
ORDER BY monthly_earned DESC
LIMIT 20;"""
for line in q2.strip().split('\n'):
    code_block(doc, line)

heading3(doc, "Q3: Atomic Points Deduction (Transaction)")
q3 = """START TRANSACTION;
  SELECT points_balance INTO @balance FROM users WHERE user_id = :userId FOR UPDATE;
  IF @balance >= 5 THEN
    UPDATE users SET points_balance = points_balance - 5 WHERE user_id = :userId;
    INSERT INTO point_records (user_id, resource_id, action_type, points_delta, balance_after)
    VALUES (:userId, :resourceId, 'SPEND_DOWNLOAD', -5, @balance - 5);
    UPDATE resources SET download_count = download_count + 1 WHERE resource_id = :resourceId;
    INSERT INTO downloads (resource_id, user_id) VALUES (:resourceId, :userId);
    SELECT 'SUCCESS' AS result;
  ELSE
    SELECT 'INSUFFICIENT_BALANCE' AS result;
  END IF;
COMMIT;"""
for line in q3.strip().split('\n'):
    code_block(doc, line)

# ===================== SECTION 9: INTERFACE DESIGN =====================
doc.add_paragraph()
heading1(doc, "9. Interface Design Specifications")
body(doc, "This section specifies the key UI screens and their data contract with the backend API. Figma prototypes are developed based on these specifications.")

iface_headers = ["Screen Name", "Key UI Elements", "API Endpoint(s)", "Input Validation Rules"]
iface_rows = [
    ["Search Page", "Search bar, Filter panel (Course Code dropdown, Year dropdown, Type radio buttons, Min Rating slider), Results list (cards with preview, rating, download count, Download button)", "GET /api/resources/search?keyword=&courseCode=&year=&type=&minRating=&page=", "keyword: 1-100 chars; courseCode: regex [A-Z]{4}\\d{5}; year: 2018-2030; stars: 1-5"],
    ["Resource Detail Page", "Title, metadata tags, description, preview/thumbnail, Download button, Rating widget (1-5 stars + comment textarea), Related Resources carousel", "GET /api/resources/{id}  |  POST /api/ratings  |  POST /api/downloads/{id}", "comment: max 500 chars; stars: integer 1-5"],
    ["Upload Page", "File dropzone (max 50MB), Title input, Course Code input, Year select, Type select, Description textarea, Tags multi-select, Submit button", "POST /api/resources/upload (multipart/form-data)", "title: 5-200 chars; courseCode: required; min 2 tags; file: PDF/DOCX/PPTX/IMAGE only"],
    ["Points Dashboard", "Points balance (large display), Monthly earned/spent summary, Transaction history table (paginated), Redemption options cards, Leaderboard preview", "GET /api/users/{id}/points  |  GET /api/points/history  |  POST /api/redemptions", "redemption optionId: valid enum value"],
    ["My Profile Page", "Avatar, Username, Student ID, Upload count, Download count, Uploaded resources grid", "GET /api/users/{id}/profile  |  GET /api/users/{id}/resources", "—"],
    ["Admin Review Queue", "Pending resources table (title, uploader, upload date), Preview pane, Approve/Reject buttons, Rejection reason input", "GET /api/admin/queue  |  PATCH /api/admin/resources/{id}/review", "rejectionReason: required if rejecting; max 500 chars"],
]
add_table(doc, iface_headers, iface_rows, [1.3, 2.2, 1.8, 1.7])

# ===================== SECTION 10: SECURITY DESIGN =====================
doc.add_paragraph()
heading1(doc, "10. Security Design")
sec_headers = ["Threat", "Risk Level", "Countermeasure", "Implementation Detail"]
sec_rows = [
    ["Unauthorized access to resources", "High", "Session-based authentication with JWT tokens", "JWT with 24h expiry; refresh tokens stored HttpOnly cookies; all API routes require valid token except /search (guest mode)"],
    ["Password breach", "High", "bcrypt password hashing", "bcrypt with cost factor 12; never store plaintext; salt auto-generated per user"],
    ["SQL Injection", "High", "Parameterized queries only", "All database calls use prepared statements with bound parameters; no string concatenation in SQL"],
    ["File upload exploits (malware)", "High", "Server-side MIME type validation + virus scan hook", "Reject files whose MIME type does not match extension; filenames sanitized; files stored outside web root"],
    ["XSS (cross-site scripting)", "Medium", "Output encoding + Content Security Policy header", "All user-generated content HTML-escaped on render; CSP header blocks inline scripts"],
    ["CSRF", "Medium", "CSRF tokens on all state-changing forms", "Double-Submit Cookie pattern; SameSite=Strict on session cookies"],
    ["Points manipulation (race condition)", "High", "Database-level locking for all balance updates", "SELECT ... FOR UPDATE in transaction; CHECK constraint balance >= 0; see Q3 in Section 8.3"],
    ["Privacy (PII exposure)", "High", "Minimum data collection; student IDs masked in API responses", "Student IDs never returned in public API responses; only username shown; compliant with PIPL"],
    ["Denial of Service", "Medium", "Rate limiting on API gateway", "Max 100 requests/min per IP; max 10 uploads/day per user; max 50 downloads/day per user"],
]
add_table(doc, sec_headers, sec_rows, [1.5, 0.9, 1.8, 2.8])

# ===================== SECTION 11: DECISION LOG =====================
doc.add_paragraph()
heading1(doc, "11. Design Decision Log")
body(doc, "Key architectural decisions and their rationale:")

dec_headers = ["Decision ID", "Decision", "Rationale", "Alternatives Considered", "Date"]
dec_rows = [
    ["DD-01", "Use Three-Tier Architecture", "Aligns with Tilley 12e course content; supports independent layer development; familiar to team", "Microservices (too complex for 3-person team within semester)", "Mar 20, 2026"],
    ["DD-02", "MySQL over NoSQL for data storage", "Course requirement (Git + MySQL); structured relational data suits the domain; team has MySQL experience", "MongoDB (flexible schema not needed; team unfamiliar)", "Mar 20, 2026"],
    ["DD-03", "Composite relevance score (40/30/30)", "Balances freshness (match accuracy), popularity (downloads), and quality (rating); simple to tune", "Pure full-text relevance ranking (ignores community quality signals)", "Apr 5, 2026"],
    ["DD-04", "5 pts per download cost", "Creates sustainable economy where active sharers always have credits; new users get 3 free daily downloads to lower entry barrier", "Free downloads (no economy); per-file pricing (too transactional)", "Apr 5, 2026"],
    ["DD-05", "Admin approval gate for uploads", "Prevents low-quality/copyrighted content from polluting search results; manageable queue size for campus-scale deployment", "Automatic publish + post-hoc flagging (risks quality at launch)", "Apr 10, 2026"],
    ["DD-06", "Atomic transaction for point deduction", "Prevents race conditions that could allow users to double-spend points", "Application-level locking (unreliable under concurrency)", "Apr 15, 2026"],
]
add_table(doc, dec_headers, dec_rows, [0.7, 1.5, 2.0, 1.7, 0.8])

# Footer
doc.add_paragraph()
footer_p = doc.add_paragraph("End of System Design Document  |  Version 1.0  |  May 2, 2026  |  System Analysis and Design — MUST")
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.runs[0].font.size = Pt(9)
footer_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save('/Users/yuxianglian/Downloads/SAD_Project/System_Design_Document.docx')
print("Done: System_Design_Document.docx")
