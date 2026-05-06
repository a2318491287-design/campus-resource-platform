"""
Patch PROJECT_CHARTER.docx — insert explicit feature-composition labeling
to comply with the BBAZ16604 requirement: "3-4 students: 1 feature
improvement and 1 new feature".

Inserts a new "1.1 Feature Composition" subsection after paragraph 11
(the "two core modules" sentence).
"""
from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy
from docx.oxml.ns import qn
from lxml import etree

CHARTER = '/Users/yuxianglian/Documents/系统分析与设计/SAD_Project/PROJECT_CHARTER.docx'

doc = Document(CHARTER)

# Find the "two core modules" paragraph
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'two core modules' in p.text:
        target_idx = i
        break

if target_idx is None:
    raise RuntimeError("Could not find target paragraph")

# Reference paragraph for insertion (we'll insert AFTER target_idx, before the empty line)
target_p = doc.paragraphs[target_idx]
target_element = target_p._element

# Build new paragraphs (heading + body x3) using existing styling as template
def new_paragraph(text, *, bold=False, size=11, color=None, after_element=None):
    """Insert a new paragraph after the given element."""
    new_p = deepcopy(target_p._element)
    # Strip existing runs
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # Strip existing pPr to avoid weird inheritance, then re-add empty
    pPr = new_p.find(qn('w:pPr'))
    if pPr is not None:
        new_p.remove(pPr)
    # Add a new run with the text
    r = etree.SubElement(new_p, qn('w:r'))
    # rPr
    rPr = etree.SubElement(r, qn('w:rPr'))
    if bold:
        b = etree.SubElement(rPr, qn('w:b'))
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(size * 2))  # docx uses half-points
    if color:
        c = etree.SubElement(rPr, qn('w:color'))
        c.set(qn('w:val'), color)
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    # text
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    after_element.addnext(new_p)
    return new_p

# Content in forward (visible) order — each new paragraph chains after the previous one
content = [
    ("1.1 Feature Composition (Course Requirement Mapping)",
     True, 13, "1F497D"),
    ("", False, 11, None),
    ("Per the BBAZ16604 course requirement that 3-student teams deliver \"1 feature improvement + 1 new feature\", the two core modules are formally classified as follows:",
     False, 11, None),
    ("", False, 11, None),
    ("Feature 1 — Improvement: Optimized Precise Retrieval Module",
     True, 11, "1F497D"),
    ("This feature improves on the existing fragmented academic resource discovery experience that students currently rely on (WeChat group searches, paid third-party platforms, peer-to-peer messaging from upperclassmen). It introduces structured multi-dimensional filtering by Course Code, Academic Year, Resource Type, and minimum Rating threshold; a composite relevance ranking formula (40% match accuracy + 30% download popularity + 30% average rating); preview snippets to reduce mistaken downloads; and personal search history. The improvement target is a 60% gain in retrieval efficiency vs current channels (per Section 2 objectives).",
     False, 11, None),
    ("", False, 11, None),
    ("Feature 2 — New Feature: Points-Based Incentive System",
     True, 11, "1F497D"),
    ("This is a new feature not present on any existing campus resource platform used by MUST students. It introduces a points economy with five action types (upload approval +10, download received +2, rating received +1, download cost -5, daily free download tier +0); two redemption options (50 pts → 10 extra downloads; 100 pts → 7-day resource pinning); a monthly campus leaderboard ranking the top 20 contributors; and a complete points transaction audit trail with running balance. The new-feature target is an 80% increase in voluntary sharing willingness (per Section 2 objectives).",
     False, 11, None),
    ("", False, 11, None),
]

# Chain insertions: each new paragraph goes after the previously inserted one
last_inserted = target_element
for text, bold, size, color in content:
    last_inserted = new_paragraph(text, bold=bold, size=size, color=color, after_element=last_inserted)

doc.save(CHARTER)
print(f"✓ Patched {CHARTER}")
print(f"  Added '1.1 Feature Composition' subsection with explicit improvement/new-feature labels")
