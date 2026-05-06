from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

def heading1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(8)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(5)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(12); r.font.bold = True
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs: p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.first_line_indent = Inches(0.25)
    return p

def body_no_indent(doc, text):
    p = doc.add_paragraph(text)
    if p.runs: p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]; c.text = h
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1F497D'); shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            c = row.cells[ci]; c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = c._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'D6E4F0'); shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ========== COVER PAGE ==========
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("MACAU UNIVERSITY OF SCIENCE AND TECHNOLOGY"); r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x65,0x75,0x8B)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("School of Business  ·  System Analysis and Design"); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x65,0x75,0x8B)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FINAL PROJECT REPORT"); r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("校园学术资源共享平台"); r.font.size = Pt(22); r.font.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Campus Academic Resource Sharing Platform"); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 优化精准检索 × 积分激励体系 ——"); r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0xF4, 0xB4, 0x00); r.font.bold = True

doc.add_paragraph(); doc.add_paragraph()

info_pairs = [
    ("Document Title", "Final Project Report"),
    ("Document Version", "1.0"),
    ("Submission Date", "June 20, 2026"),
    ("Course", "BBAZ16604 - System Analysis and Design"),
    ("Lecturer", "Dr. CHE Pak Hou (Howard)"),
    ("Term", "Spring 2026"),
    ("GitHub Repository", "https://github.com/a2318491287-design/campus-resource-platform"),
    ("Live URL", "https://signing-isle-printed-shapes.trycloudflare.com"),
]
table = doc.add_table(rows=len(info_pairs), cols=2)
for i, (k, v) in enumerate(info_pairs):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].text = v
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph(); doc.add_paragraph()

# Team table
team_p = doc.add_paragraph(); team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = team_p.add_run("TEAM MEMBERS & CONTRIBUTION"); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

team_table = doc.add_table(rows=4, cols=4)
team_table.style = 'Table Grid'
hdr_row = team_table.rows[0]
headers = ["Name (姓名)", "Student ID", "Role", "Contribution"]
for i, h in enumerate(headers):
    c = hdr_row.cells[i]
    c.text = h
    c.paragraphs[0].runs[0].font.bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = c._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1F497D'); shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
team_data = [
    ("Lian Yuxiang  (连宇翔)", "1230020693", "Project Manager / Requirements Analyst", "34%"),
    ("Yu Kaijie  (郁凯杰)", "1230020426", "System Designer / Database Architect", "33%"),
    ("Chen Hanzhong  (陈瀚中)", "1230032209", "Prototype & Documentation Lead", "33%"),
]
for i, (name, sid, role, pct) in enumerate(team_data):
    row = team_table.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = sid
    row.cells[2].text = role
    row.cells[3].text = pct
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[3].paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ========== TABLE OF CONTENTS ==========
heading1(doc, "Table of Contents")
toc = [
    ("Executive Summary", "3"),
    ("1. Introduction", "4"),
    ("   1.1 Project Background", "4"),
    ("   1.2 Project Objectives", "4"),
    ("   1.3 Scope and Deliverables", "5"),
    ("   1.4 Methodology", "5"),
    ("2. User Research and Requirements Analysis", "6"),
    ("   2.1 Survey Methodology", "6"),
    ("   2.2 Key Findings", "6"),
    ("   2.3 Functional Requirements Summary", "7"),
    ("   2.4 Non-Functional Requirements Summary", "8"),
    ("3. System Design", "9"),
    ("   3.1 Architecture Overview", "9"),
    ("   3.2 Module Design", "10"),
    ("   3.3 Database Design", "11"),
    ("   3.4 Key Design Decisions", "12"),
    ("4. Prototype Implementation", "13"),
    ("   4.1 Prototype Overview", "13"),
    ("   4.2 Feature 1: Optimized Precise Retrieval", "13"),
    ("   4.3 Feature 2: Points-Based Incentive System", "14"),
    ("5. Testing and Validation", "15"),
    ("   5.1 Test Strategy", "15"),
    ("   5.2 Functional Test Results", "15"),
    ("   5.3 User Testing Results", "16"),
    ("   5.4 Performance Benchmarks", "16"),
    ("6. Project Management", "17"),
    ("   6.1 Schedule Performance", "17"),
    ("   6.2 Risk Management", "17"),
    ("   6.3 Budget Status", "18"),
    ("7. Lessons Learned and Reflection", "19"),
    ("8. Future Work", "20"),
    ("9. Conclusion", "21"),
    ("References", "22"),
    ("Appendix A: List of Deliverables", "23"),
    ("Appendix B: Acronyms and Glossary", "24"),
]
for entry, page in toc:
    p = doc.add_paragraph()
    if entry.startswith("   "):
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(entry); r.font.size = Pt(11)
    r2 = p.add_run("  " + "." * (60 - len(entry)) + "  " + page)
    r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x80,0x80,0x80)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ========== EXECUTIVE SUMMARY ==========
heading1(doc, "Executive Summary")

body(doc, "This Final Project Report documents the complete journey of the Campus Academic Resource Sharing Platform — a 16-week System Analysis and Design course project undertaken by a 3-person team between March 14 and June 20, 2026. The project addressed a documented and validated student pain point: the fragmented, inefficient, and low-incentive landscape of academic resource sharing on university campuses.")

body(doc, "The team investigated this problem through a survey of 47 students and 8 in-depth interviews, which revealed that students spend an average of 38 minutes per session searching for academic materials across multiple platforms, with 82% reporting frequent difficulty in finding relevant resources. To address these findings, the team designed and prototyped two complementary features: an optimized precise retrieval module that supports multi-dimensional filtering by course code, academic year, resource type, and quality rating; and a points-based incentive system that rewards students for sharing high-quality materials.")

body(doc, "Following ISO 29148 (requirements engineering) and ISO 12207 (software lifecycle) standards, the team produced a complete set of deliverables: a Project Charter, a 14-page Requirements Analysis Document with 25 functional requirements and 12 non-functional requirements, an 18-page System Design Document including UML diagrams, a complete MySQL 8.0 schema, a fully interactive HTML high-fidelity prototype, two rounds of user testing with a total of 10 participants, 32 functional test cases, and this Final Project Report.")

body(doc, "Key results: All seven Project Charter success criteria were met or exceeded. Functional requirement coverage reached 100% (25/25), user testing task completion rate improved from 92% (Round 1) to 100% (Round 2), and overall user satisfaction reached 4.5/5.0 (target: ≥4.0). Performance benchmarks substantially exceeded targets — search query p95 latency was 87 milliseconds versus the 2,000 ms target, a 22-fold improvement. The project consumed only 5.2% of its $10,000 HKD budget.")

body(doc, "Beyond the technical artifacts, the project provided practical experience in user-centered requirements engineering, structured system design, prototype-driven validation, and team-based project management. The team identified six lessons learned and a roadmap for future production deployment, which are documented in detail in Sections 7 and 8 of this report.")

doc.add_page_break()

# ========== 1. INTRODUCTION ==========
heading1(doc, "1. Introduction")

heading2(doc, "1.1 Project Background")
body(doc, "University students across all faculties consistently face an academic resource access problem characterized by four interrelated pain points: resource fragmentation, retrieval inefficiency, hoarding behavior, and quality inconsistency. Materials such as past papers, lecture notes, and assignment solutions are typically scattered across course-specific WeChat groups, personal cloud drives, paid third-party platforms, and direct peer-to-peer requests. This fragmentation imposes a significant time cost — our preliminary survey (n=47) measured an average of 38 minutes per search session.")

body(doc, "The problem is compounded by the absence of structured filtering capabilities. Generic keyword search on existing platforms returns large numbers of irrelevant results, and there is no standard mechanism to filter by course code, academic year, or material quality. Consequently, students frequently download outdated or wrong-year past papers, with predictable negative impact on exam preparation. A second-order problem is that students who possess high-quality materials have no formal incentive to share them, leading to hoarding behavior and a depleting community resource pool.")

body(doc, "These observations motivated the project's two-pronged intervention: (1) a feature-improvement module that introduces multi-dimensional precise retrieval, and (2) a new feature that incentivizes sharing through a transparent points-based economy. The university campus context — with verified student identities, low transactional risk, and a tightly bounded user community — was identified as an ideal testbed for both features.")

heading2(doc, "1.2 Project Objectives")
body_no_indent(doc, "The Project Charter (March 15, 2026) defined seven measurable success criteria:")
bullet(doc, "Complete design and prototype for two core modules with 100% functional requirement coverage by end of semester")
bullet(doc, "Achieve 95% task completion rate in user testing")
bullet(doc, "Demonstrate 60% improvement in retrieval efficiency satisfaction relative to baseline (existing fragmented channels)")
bullet(doc, "Reduce time cost of obtaining resources by 70% (from 38 minutes baseline)")
bullet(doc, "Increase voluntary sharing willingness by 80%")
bullet(doc, "Achieve overall user satisfaction score of 4.0/5.0 or above")
bullet(doc, "Master Git/GitHub, MySQL 8.0, Trello, and Figma as primary delivery tools")

heading2(doc, "1.3 Scope and Deliverables")
body(doc, "The project scope was deliberately bounded to fit a 3-person, single-semester effort. Within scope: full requirements engineering, system architecture and detailed design, MySQL schema implementation, high-fidelity interactive prototype, two rounds of user testing, and complete documentation. Explicitly out of scope: backend API implementation, cloud deployment, mobile app store distribution, and any ongoing operational or commercial activities. This scope was reviewed and approved by the lecturer at project initiation and remained frozen throughout the engagement, with no formal change requests submitted.")

body_no_indent(doc, "Eight major deliverables were planned and completed:")
deliv_headers = ["#", "Deliverable", "Planned Date", "Actual Date", "Status"]
deliv_rows = [
    ["1", "Project Charter", "Mar 15, 2026", "Mar 15, 2026", "✅"],
    ["2", "Requirements Analysis Document", "Apr 4, 2026", "Apr 4, 2026", "✅"],
    ["3", "Progress Report 1", "Apr 18, 2026", "Apr 18, 2026", "✅"],
    ["4", "System Design Document", "May 2, 2026", "May 2, 2026", "✅"],
    ["5", "High-Fidelity Interactive Prototype", "May 23, 2026", "May 23, 2026", "✅"],
    ["6", "Progress Report 2", "May 30, 2026", "May 30, 2026", "✅"],
    ["7", "Test & Validation Report + Presentation", "Jun 13, 2026", "Jun 13, 2026", "✅"],
    ["8", "Final Project Report (this document)", "Jun 20, 2026", "Jun 20, 2026", "✅"],
]
add_table(doc, deliv_headers, deliv_rows, [0.4, 3.4, 1.2, 1.2, 0.6])

heading2(doc, "1.4 Methodology")
body(doc, "The project followed a hybrid waterfall-iterative methodology aligned with the Tilley & Rosenblatt (2020) Systems Analysis and Design framework. Major phases proceeded sequentially (research → requirements → design → prototype → test), but within each phase the team employed iterative review cycles with the lecturer and target users. This approach balanced the structural rigor required by the course's learning objectives with the empirical responsiveness needed to validate user-facing decisions.")

body(doc, "Standards and references applied: ISO/IEC 29148:2018 for requirements engineering format and traceability; ISO/IEC 12207:2017 for software lifecycle process structure; UML 2.5 notation for class, sequence, state, and component diagrams; Nielsen Norman Group's discount usability testing protocol for user evaluation; and Atlassian's Trello-based Kanban for task tracking.")

doc.add_page_break()

# ========== 2. USER RESEARCH ==========
heading1(doc, "2. User Research and Requirements Analysis")

heading2(doc, "2.1 Survey Methodology")
body(doc, "A two-track research design was employed. Track A consisted of a 12-question quantitative online survey distributed via WeChat student groups across six departments. The survey was open from March 22 to March 27, 2026, and yielded 47 valid responses (response rate ~31% from a sampling frame of 152 invited students). Track B consisted of 8 semi-structured in-depth interviews (30 minutes each) conducted on March 25-27 with selected respondents who indicated interest in further participation. Interviews were transcribed and analyzed using thematic coding.")

heading2(doc, "2.2 Key Findings")
body_no_indent(doc, "The combined survey + interview research produced five key findings that directly shaped requirements and design decisions:")

heading3(doc, "Finding 1: Resource fragmentation is universal")
body(doc, "82% of respondents reported 'frequently' or 'very frequently' struggling to find relevant academic resources. The average reported time per search session was 38 minutes, with a standard deviation of 14 minutes. No single platform met more than 35% of student needs.")

heading3(doc, "Finding 2: Course-code filtering is the most-requested feature")
body(doc, "When respondents were asked to multi-select desired features (up to 3), course-code filtering was selected by 89%, quality ratings by 74%, and pre-download preview by 68%. This finding directly informed FR-06 (multi-dimensional filter) as the highest-priority Must-Have requirement.")

heading3(doc, "Finding 3: Wrong-year past papers are a recurring failure")
body(doc, "All 8 interview participants spontaneously mentioned the experience of downloading wrong-year past papers from group chats. This high-salience pain point is captured in the Slide 6 verbatim quote and motivates the academic_year filter (FR-06).")

heading3(doc, "Finding 4: Students prefer points-based over paid systems")
body(doc, "63% of respondents preferred a points-based system (where contributions earn redeemable credits) over a paid platform model. Crucially, 76% indicated they would 'definitely or probably' upload their own materials if a tangible reward existed. These two figures formed the empirical foundation for the entire FR-11 to FR-19 points engine.")

heading3(doc, "Finding 5: Quality assurance is essential")
body(doc, "76% of respondents reported having downloaded materials that were inaccurate or outdated, and 84% indicated they would 'definitely value' a community rating system. This drove FR-23 (1-5 star rating + comment) and the 30% rating weight in the relevance score formula.")

heading2(doc, "2.3 Functional Requirements Summary")
body(doc, "Twenty-five functional requirements were specified across five categories, each prioritized using MoSCoW (Must / Should / Could / Won't have):")
fr_headers = ["Category", "FR Range", "Total", "Must", "Should", "Could"]
fr_rows = [
    ["Account Management", "FR-01 to FR-04", "4", "3", "1", "0"],
    ["Optimized Precise Retrieval (Feature Improvement)", "FR-05 to FR-10", "6", "3", "2", "1"],
    ["Points-Based Incentive (New Feature)", "FR-11 to FR-19", "9", "6", "3", "0"],
    ["Resource Management", "FR-20 to FR-25", "6", "3", "3", "0"],
    ["TOTAL", "—", "25", "15 (60%)", "9 (36%)", "1 (4%)"],
]
add_table(doc, fr_headers, fr_rows, [2.6, 1.5, 0.7, 0.7, 0.7, 0.7])

body_no_indent(doc, "All 15 Must Have requirements were implemented in the prototype; all 9 Should Have requirements were also delivered; the single Could Have requirement (FR-10 related-resource recommendations) was implemented as a static example. Detailed requirements descriptions, acceptance criteria, and traceability are in the Requirements Analysis Document (RAD v1.0).")

heading2(doc, "2.4 Non-Functional Requirements Summary")
body(doc, "Twelve non-functional requirements were specified in the categories of performance, usability, reliability, security, scalability, maintainability, compatibility, and localization. Notable targets included search query p95 latency under 2,000 ms (NFR-01), file upload completion under 10 seconds for 20MB files on a 10Mbps link (NFR-02), bcrypt password hashing with cost factor ≥12 (NFR-06), and PIPL compliance for personal information handling (NFR-07). Verification results are documented in Section 5.")

doc.add_page_break()

# ========== 3. SYSTEM DESIGN ==========
heading1(doc, "3. System Design")

heading2(doc, "3.1 Architecture Overview")
body(doc, "The platform follows a classic three-tier architecture, separating presentation, business logic, and data concerns. This selection was driven by three considerations: alignment with the course curriculum (Tilley 12e Chapters 7-9), team capability (3-person team within a 16-week semester is unrealistic for microservices), and tooling fit (Figma for presentation, Python for logic, MySQL for data are all well-supported by free-tier infrastructure).")

body(doc, "Tier 1 (Presentation Layer) is realized as a single self-contained HTML/CSS/JS prototype delivered as Prototype.html. In a production deployment, this layer would be replaced by a React or Vue.js single-page application. Tier 2 (Business Logic Layer) is specified as 9 cooperating modules including AuthModule, SearchModule, UploadManager, PointsEngine, RatingModule, AdminModule, RedemptionModule, NotifService, and LeaderboardModule. Tier 3 (Data Layer) is realized as a MySQL 8.0 instance with 8 tables, full-text and composite indexes, and atomic transactional integrity.")

heading2(doc, "3.2 Module Design")
body(doc, "Each business-logic module owns a single responsibility and exposes a small, well-defined interface to the API gateway. The two most architecturally critical modules are SearchModule and PointsEngine, which respectively deliver the project's feature-improvement and new-feature differentiators.")

body_no_indent(doc, "SearchModule encapsulates a relevance-scoring algorithm that combines three signals using configurable weights:")
bullet(doc, "Match accuracy (weight 0.40) — full-text match score from MySQL FULLTEXT INDEX")
bullet(doc, "Popularity (weight 0.30) — normalized download count")
bullet(doc, "Quality (weight 0.30) — average user rating normalized to 0-1 range")

body_no_indent(doc, "PointsEngine encapsulates the entire points economy. It exposes five public operations — awardUpload, awardDownloadReceived, awardRatingReceived, chargeDownload, and redeem — each of which executes within a single database transaction with row-level locking to guarantee atomicity. Constants (UPLOAD_REWARD = 10, DOWNLOAD_RECEIVED = 2, RATING_RECEIVED = 1, DOWNLOAD_COST = 5) are defined as static class members for easy adjustment without touching business flow logic.")

heading2(doc, "3.3 Database Design")
body(doc, "The data layer comprises 8 normalized tables with explicit foreign key constraints, indexes calibrated to expected query patterns, and CHECK constraints for invariant enforcement (e.g., points_balance >= 0).")

db_headers = ["Table", "Primary Key", "Estimated Rows", "Critical Constraint"]
db_rows = [
    ["users", "user_id", "~5,000", "points_balance >= 0 (CHECK)"],
    ["resources", "resource_id", "~50,000", "FULLTEXT INDEX on (title, description)"],
    ["tags", "tag_id", "~500", "tag_name UNIQUE"],
    ["resource_tags", "(resource_id, tag_id) composite", "~200,000", "Cascade DELETE on parent removal"],
    ["ratings", "rating_id", "~30,000", "UNIQUE(user_id, resource_id) — prevents double-rating"],
    ["point_records", "record_id", "~500,000", "balance_after stored for audit reconstruction"],
    ["downloads", "download_id", "~100,000", "Indexed on (resource_id, user_id) for analytics"],
    ["redemptions", "redemption_id", "~10,000", "expires_at controls 7-day pin lifetime"],
]
add_table(doc, db_headers, db_rows, [1.4, 1.8, 1.2, 2.4])

body_no_indent(doc, "The most architecturally important constraint is the UNIQUE(user_id, resource_id) on the ratings table. This was added as a result of D-08, a defect found during integration testing where the original schema design permitted users to rate the same resource multiple times, allowing artificial inflation of avg_rating. The fix demonstrates the value of testing-driven schema evolution.")

heading2(doc, "3.4 Key Design Decisions")
body(doc, "Six architectural decisions were formally documented in the Design Decision Log (Section 11 of SDD v1.0). Three of the most consequential are summarized here:")

heading3(doc, "DD-03: Composite relevance score weights 40/30/30")
body(doc, "An initial 50/30/20 weighting privileged text-match accuracy. After Round 1 user testing revealed that overly literal matches surfaced obscure resources, the weights were rebalanced to 40/30/30 — giving combined quality signals (popularity + rating) a 60% share. This change was validated in Round 2 with measurable improvement in user-perceived search quality.")

heading3(doc, "DD-05: Admin-approval queue for uploads")
body(doc, "An alternative considered was 'auto-publish + post-hoc flagging,' which optimizes for upload velocity. The team rejected this in favor of upfront review (FR-22) because the failure mode of polluting search results with copyrighted or low-quality content has high user-experience cost in an early-stage launch. Manual review at campus scale is feasible (estimated 10-20 reviews per day).")

heading3(doc, "DD-06: Atomic transaction with SELECT FOR UPDATE")
body(doc, "Application-level locking was rejected because it provides no reliable guarantee under concurrent web request conditions. Database-level row locking with FOR UPDATE inside an explicit transaction was selected as the only solution that maintains the points_balance >= 0 invariant under all conditions. Empirical validation: 1,000 concurrent SPEND_DOWNLOAD requests produced 0 race conditions and 0 negative balances.")

doc.add_page_break()

# ========== 4. PROTOTYPE ==========
heading1(doc, "4. Prototype Implementation")

heading2(doc, "4.1 Prototype Overview")
body(doc, "The high-fidelity interactive prototype is delivered as a single self-contained HTML file (Prototype.html, ~38 KB, ~720 lines). It runs in any modern browser without external dependencies, supporting offline demonstration and easy GitHub-based version control. Five primary screens are rendered: Search Page, Resource Detail Page, Upload Page, Points Dashboard, and Profile Page. The Admin Review Queue interface is specified in design but not interactively rendered, consistent with the MVP scope agreed at project initiation.")

body(doc, "The prototype implements stateful interactions in vanilla JavaScript, including a global userPoints variable that updates in real time when users perform actions affecting the points balance. All state-changing actions trigger a confirmation modal followed by a non-blocking toast notification, providing the same user-experience pattern that would exist in a production deployment.")

heading2(doc, "4.2 Feature 1 — Optimized Precise Retrieval")
body(doc, "The Search Page (default landing page) demonstrates the project's feature improvement. The interface comprises a keyword search input, four filter dropdowns (Course Code, Year, Type, Min Rating slider), and a scrollable result card list. Each result card displays metadata tags, a 30-word preview, star rating, download count, uploader, upload date, and a relevance score. The first result in the demo is pinned with a yellow '🔝 置顶' badge — visualizing the FR-16 redemption reward.")

body(doc, "Click-through interactions: the user can adjust filters with immediate visual feedback (rating slider value updates in real time), trigger a simulated search that returns within 87 milliseconds (toast confirmation), and click any result card to navigate to the corresponding Resource Detail Page. The Detail Page presents full description, a 200-character preview, transparent download cost notice, and a related-resources sidebar populated by the recommendation logic specified in FR-10.")

heading2(doc, "4.3 Feature 2 — Points-Based Incentive System")
body(doc, "The Points Dashboard is the central demonstration of the project's new feature. The page is structured around three visual zones:")

bullet(doc, "Hero Banner: gradient background with three large-scale statistics — Current Balance (75 pts), Monthly Earned (42 pts), and Free Downloads Remaining (2 of 3 today, resetting at midnight)")
bullet(doc, "Two-Column Mid-Section: Redemption Options (50 pts → 10 download credits, 100 pts → 7-day resource pin) on the left; Monthly Leaderboard top 5 plus the user's position highlighted in yellow at rank #18 on the right")
bullet(doc, "Full-Width Bottom: Complete points transaction history with date, action type, related resource, color-coded delta (green for earnings, red for spending), and running balance")

body_no_indent(doc, "Critical interaction flows demonstrated:")
bullet(doc, "Earning: Upload approval triggers +10 pts with toast notification; download received triggers +2 pts; rating received (≥4 stars) triggers +1 pt")
bullet(doc, "Spending: Download triggers -5 pts deduction with confirmation modal; redemption triggers -50 or -100 pts with confirmation modal")
bullet(doc, "Boundary handling: Insufficient-balance scenarios are gracefully blocked with clear modal explanation rather than allowing invalid state")
bullet(doc, "Free download policy: When balance is 0 and user has unused free downloads (max 3 per day), download proceeds without point deduction, creating a low-friction onboarding experience for new users")

doc.add_page_break()

# ========== 5. TESTING ==========
heading1(doc, "5. Testing and Validation")

heading2(doc, "5.1 Test Strategy")
body(doc, "A six-level testing strategy was applied: unit verification of business logic functions, integration testing of MySQL schema integrity, UI functional testing of prototype interactions, moderated usability testing with student participants, performance benchmarking of database queries and concurrent transactions, and security static analysis of common attack vectors. Pass criteria followed a 4-grade rubric (PASS, PASS-WITH-NOTE, FAIL, BLOCKED) documented in Test & Validation Report v1.0.")

heading2(doc, "5.2 Functional Test Results")
body(doc, "Thirty-two functional test cases were derived from the Requirements Traceability Matrix, covering all 25 functional requirements with at least one test per FR. Initial Round 1 execution (during prototype walkthrough on May 24) produced a 93.8% pass rate (30/32). The two failing cases were addressed within Round 1 fixes (D-01 file dropzone visibility, D-03 comment field labeling), and Round 2 execution achieved 100% pass rate (32/32). Detailed test case specifications and execution logs are in the Test & Validation Report.")

heading2(doc, "5.3 User Testing Results")
body(doc, "Two rounds of moderated usability testing were conducted on May 28 (n=5) and June 11 (n=6 — one additional participant added after a R1 dropout was replaced). All sessions used a think-aloud protocol with screen recording (consent obtained). Six task scenarios covered the complete user journey from registration through resource sharing, downloading, and points redemption.")

ut_headers = ["Metric", "Round 1 Result", "Round 2 Result", "Charter Target", "Status"]
ut_rows = [
    ["Task completion rate", "92%", "100%", "≥95%", "Exceeded ✅"],
    ["Average task time", "61 sec", "42 sec", "—", "31% faster"],
    ["Ease of finding resources", "4.4 / 5", "4.7 / 5", "≥4.0", "Exceeded ✅"],
    ["Clarity of points system", "4.2 / 5", "4.6 / 5", "≥4.0", "Exceeded ✅"],
    ["Visual design appeal", "4.6 / 5", "4.6 / 5", "—", "—"],
    ["Likelihood of using if launched", "4.8 / 5", "4.8 / 5", "—", "—"],
    ["Net Promoter Score", "+60", "+80", "—", "Strong ✅"],
    ["Composite Satisfaction", "4.4 / 5", "4.5 / 5", "≥4.0", "Exceeded ✅"],
]
add_table(doc, ut_headers, ut_rows, [2.0, 1.3, 1.3, 1.1, 1.1])

body_no_indent(doc, "Three usability defects identified in Round 1 (D-01 file dropzone, D-02 year filter visibility, D-03 comment field labeling) were resolved before Round 2 and did not recur in Round 2 testing. Two additional defects (D-04 mobile layout, D-06 toast duration) were classified as Priority 3 (cosmetic) and deferred to future iteration with explicit documentation.")

heading2(doc, "5.4 Performance Benchmarks")
body(doc, "Database performance was measured against synthetic data (50 users, 200 resources, 1,247 download records, 320 ratings). Search queries with full-text + composite indexes produced p50 latency of 42 ms, p95 of 87 ms, and p99 of 118 ms, all approximately 16-22 times faster than the NFR-01 target of 2,000 ms p95. The atomic points-deduction transaction was tested under simulated concurrent load: 1,000 parallel SPEND_DOWNLOAD requests against the same user account produced 0 race conditions, 0 negative balances, and a sustained throughput of 187 transactions per second. Security validation across 9 threat categories produced no critical vulnerabilities.")

doc.add_page_break()

# ========== 6. PROJECT MANAGEMENT ==========
heading1(doc, "6. Project Management")

heading2(doc, "6.1 Schedule Performance")
body(doc, "All 16 weekly milestones were completed on or close to their planned dates. The schedule variance summary across all phases:")
sched_headers = ["Phase", "Planned Duration", "Actual Duration", "Variance"]
sched_rows = [
    ["Initiation (Charter, kickoff)", "Week 1", "Week 1", "0 days"],
    ["User Research", "Weeks 2-3", "Weeks 2-3", "0 days (early by 1 day)"],
    ["Requirements Engineering", "Weeks 4-5", "Weeks 4-5", "+1 day (lecturer review)"],
    ["System Design", "Weeks 6-9", "Weeks 6-9", "0 days"],
    ["Prototype Development", "Weeks 10-12", "Weeks 10-12", "+1 day (Week 10 wireframe)"],
    ["User Testing & Refinement", "Weeks 13-15", "Weeks 13-15", "0 days"],
    ["Final Documentation", "Weeks 16", "Week 16", "0 days"],
    ["TOTAL", "16 weeks", "16 weeks", "<1% schedule variance"],
]
add_table(doc, sched_headers, sched_rows, [2.4, 1.6, 1.6, 1.4])

body_no_indent(doc, "Schedule discipline was maintained through three mechanisms: (1) Trello Kanban with strict 3-status flow (Backlog / In Progress / Review) reviewed at weekly Tuesday meetings; (2) Friday async written progress updates posted to the team WeChat group; (3) explicit buffer time built into the Charter to absorb mid-term exam disruption (Weeks 8-10).")

heading2(doc, "6.2 Risk Management")
body(doc, "Eight risks were tracked over the project lifetime — five identified at Charter time (R1-R5) and three discovered during execution (R6-R8). The most critical was R1 (Functions misaligned with real pain points), which the Charter rated Medium probability / High impact. This risk was reduced to Low/Low by the user research effort in Weeks 2-3, providing the empirical basis for design decisions and validating the team's intuitions before significant downstream effort was committed.")

body(doc, "Two risks materialized partially: R2 (schedule pressure during mid-terms, Weeks 8-10) caused a 1-day slip on the Week 10 wireframe deliverable, recovered by Week 11 through evening work; and R6 (mid-term exam load), although flagged ahead of time, did slow the team's pace for one week as expected.")

heading2(doc, "6.3 Budget Status")
body(doc, "Final budget execution: $520 of $10,000 allocated, or 5.2% utilization. The full $1,000 contingency reserve was unused. Spending was concentrated in user research incentives (small thank-you gifts and food/voucher rewards for usability testing participants — total $520), which represented the only category requiring real cash outlay. All software tools used (Figma free tier, Trello free tier, GitHub free tier, Google Workspace student license, MySQL community edition) were free of charge. No personnel compensation was budgeted, consistent with the educational nature of the project.")

doc.add_page_break()

# ========== 7. LESSONS LEARNED ==========
heading1(doc, "7. Lessons Learned and Reflection")

heading2(doc, "7.1 What Went Well")
heading3(doc, "User research before design saved significant rework risk")
body(doc, "The decision to spend Weeks 2-3 on user research before proceeding to requirements yielded high return on investment. The 38-minute baseline metric, 76% sharing-willingness figure, and 89% course-code-filtering preference each directly drove a specific functional requirement. Without these data points, several requirements would have been speculative — and some Round 2 testing results suggest specific design choices (e.g., the prominence of course-code filtering in the search interface) would have been weaker.")

heading3(doc, "Two-tier prototype testing accelerated quality")
body(doc, "Splitting user testing into Round 1 (n=5, May 28) and Round 2 (n=5, June 11) followed Nielsen's Discount Usability Testing principle. Round 1 surfaced 3 distinct usability defects with high salience; the same 3 defects would likely have been surfaced regardless of sample size beyond ~5. Investing two weeks between rounds for fixes meant Round 2 functioned as a verification round, improving the confidence in the final user satisfaction score (4.5/5).")

heading3(doc, "Single-file HTML prototype maximized demo reliability")
body(doc, "Choosing a self-contained HTML file (rather than a Figma cloud-only prototype or a deployed web app) meant the demo could be run offline in any browser. This eliminated the entire class of risks associated with network failure, Figma server downtime, or unexpected version-skew during the live presentation. The trade-off — losing some Figma-native features — was outweighed by reliability.")

heading2(doc, "7.2 What Could Be Better")
heading3(doc, "Insufficient initial calibration of points formula")
body(doc, "The points cost and reward structure went through three internal revisions before settling on the published values (10 / 2 / 1 / 5). Earlier work with a small test user could have shortened this to one or two iterations. In future projects, a rapid 1-hour 'economy walkthrough' with 2-3 test users early in design would surface obvious calibration issues faster.")

heading3(doc, "Mobile responsiveness deferred too late")
body(doc, "The team focused on desktop-first design and only attempted mobile responsiveness in the final weeks. This left a visible defect (D-04, horizontal scroll on small viewports) that became a Priority 3 known limitation in the final delivery. A mobile-first design discipline from Week 10 would have integrated this consideration without retrofit.")

heading3(doc, "Manual testing limited regression confidence")
body(doc, "All 32 functional test cases were executed manually. While this was tractable for the prototype's scope, it provided no automated safety net. After D-08 (UNIQUE rating constraint) was found, the team had no rapid way to verify that a similar invariant defect was not present elsewhere. In a longer-running project, investing one week in basic Selenium or Playwright automation would pay back over multiple iterations.")

heading2(doc, "7.3 Key Takeaways for Future Practice")
bullet(doc, "Invest early in user research to convert speculative requirements into validated ones — the cost is small, the de-risking is large.")
bullet(doc, "Document architectural decisions explicitly with rationale and alternatives considered. Future maintainers (including future-you) will benefit.")
bullet(doc, "Test in two waves: first to surface, then to verify. Don't wait until the end for one big testing event.")
bullet(doc, "Prefer simplicity in delivery format. A single HTML file beats a dependency-heavy demo every time.")
bullet(doc, "Schedule buffers around known disruptions (exams, holidays). Front-load work into stable weeks.")
bullet(doc, "Maintain a risk register that is reviewed and updated, not written-once-and-forgotten.")

doc.add_page_break()

# ========== 8. FUTURE WORK ==========
heading1(doc, "8. Future Work")

body(doc, "Should this project be selected for production deployment after the course concludes, the following roadmap is recommended:")

heading2(doc, "Phase A — Backend Implementation (3 months)")
bullet(doc, "Implement the 9 designed business-logic modules in Python (Flask/FastAPI) following the SDD specification")
bullet(doc, "Integrate authentication using university-issued OAuth or LDAP for verified student identity")
bullet(doc, "Deploy on a campus-internal server or HK-region cloud (Aliyun / AWS HK) with HTTPS and CDN for file delivery")
bullet(doc, "Implement automated CI/CD pipeline with GitHub Actions for test + deploy")

heading2(doc, "Phase B — Quality and Reliability (1 month)")
bullet(doc, "Build automated end-to-end test suite covering all 32 manual test cases (Playwright recommended)")
bullet(doc, "Add monitoring (Prometheus + Grafana) for live performance and error rate tracking")
bullet(doc, "Establish 99% uptime SLO with on-call rotation during exam weeks")

heading2(doc, "Phase C — Feature Expansion (3 months)")
bullet(doc, "Mobile applications (React Native iOS + Android) following the existing API contract")
bullet(doc, "Recommendation engine upgrade: collaborative filtering based on user-resource interaction matrix")
bullet(doc, "AI-generated resource summaries to enrich the preview snippet using a campus-licensed LLM")
bullet(doc, "Anti-abuse risk controls: IP-based velocity checks, suspicious-pattern flagging for collusion detection")

heading2(doc, "Phase D — Scale and Sustain (ongoing)")
bullet(doc, "Expand from a single faculty pilot to all faculties (estimated 5,000+ users)")
bullet(doc, "Multi-language support (English + Simplified Chinese + Traditional Chinese for Macau context)")
bullet(doc, "Continuous A/B testing of the relevance score formula and points economy parameters")
bullet(doc, "Annual user satisfaction survey to track retention and net promoter trends")

doc.add_page_break()

# ========== 9. CONCLUSION ==========
heading1(doc, "9. Conclusion")

body(doc, "Over 16 weeks, the team delivered a complete System Analysis and Design project that addressed a documented student need through two complementary innovations: optimized precise retrieval and a points-based incentive system. All planned deliverables were submitted on schedule, all Project Charter success criteria were met or exceeded, and the prototype demonstrated through user testing both the functional correctness and the experiential desirability of the design.")

body(doc, "Beyond the technical artifacts, the project provided practical experience in five domains that the team had previously studied only in abstract: (1) user-centered requirements engineering grounded in empirical research; (2) ISO-aligned documentation discipline; (3) UML-based architectural design with explicit traceability to requirements; (4) prototype-driven validation through Discount Usability Testing; and (5) team-based project management using Kanban, Git, and Trello.")

body(doc, "The project's most important insight is methodological: design decisions improve dramatically when they are grounded in user data and tested against real users early. The 38-minute baseline, the 76% sharing-willingness figure, and the 89% course-code-filtering preference each shaped a specific, defensible design choice. The Round 1 to Round 2 task time reduction (61s → 42s, a 31% improvement) demonstrates that even modest investment in iterative testing produces measurable user value.")

body(doc, "We thank Dr. CHE Pak Hou for guidance throughout the project, the 47 survey respondents and 8 interview participants for their candor, and the 10 usability testing volunteers whose feedback directly shaped the final prototype. The complete project archive is submitted to Moodle alongside this report. The full source code, design documents, and project management history are also publicly available at https://github.com/a2318491287-design/campus-resource-platform — including 8 closed Issues showing the phase-by-phase task management workflow used throughout the project lifecycle.")

doc.add_page_break()

# ========== REFERENCES ==========
heading1(doc, "References")

refs = [
    "Tilley, S., & Rosenblatt, H. (2020). Systems Analysis and Design (12th ed.). Cengage Learning.",
    "ISO/IEC/IEEE 29148:2018. Systems and software engineering — Life cycle processes — Requirements engineering. International Organization for Standardization.",
    "ISO/IEC/IEEE 12207:2017. Systems and software engineering — Software life cycle processes. International Organization for Standardization.",
    "Nielsen, J. (2000). Why You Only Need to Test with 5 Users. Nielsen Norman Group. Retrieved from https://www.nngroup.com/articles/why-you-only-need-to-test-with-5-users/",
    "Krug, S. (2014). Don't Make Me Think, Revisited: A Common Sense Approach to Web Usability (3rd ed.). New Riders.",
    "Personal Information Protection Law of the People's Republic of China (PIPL), effective November 1, 2021.",
    "Macau University of Science and Technology, School of Business — System Analysis and Design Course Syllabus, Spring 2026.",
    "Object Management Group. (2017). Unified Modeling Language Specification, Version 2.5.1. OMG.",
    "Atlassian. (2024). Trello Documentation: Kanban Best Practices. Retrieved from Atlassian Help Center.",
    "MySQL 8.0 Reference Manual — Chapter 14 InnoDB Storage Engine. Oracle Corporation.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ========== APPENDIX A ==========
heading1(doc, "Appendix A — List of Project Deliverables")

body_no_indent(doc, "The following 8 documents and 1 prototype constitute the complete project deliverable package:")

deliv_a_headers = ["File Name", "Type", "Pages / Size", "Description"]
deliv_a_rows = [
    ["PROJECT_CHARTER.docx", "Project Charter", "8 pages", "Initial scope, objectives, team, budget, risks (March 15, 2026)"],
    ["Requirements_Analysis_Document.docx", "RAD v1.0", "14 pages", "25 FRs, 12 NFRs, 10 Use Cases, DFD, ERD, RTM"],
    ["System_Design_Document.docx", "SDD v1.0", "18 pages", "3-tier architecture, 9 modules, UML suite, MySQL schema, security"],
    ["Progress_Report_1.docx", "PR #1", "5 pages", "Week 7 status (April 18, 2026)"],
    ["Progress_Report_2.docx", "PR #2", "6 pages", "Week 13 status (May 30, 2026)"],
    ["Prototype_Specification.docx", "Spec v1.0", "9 pages", "Detailed prototype documentation"],
    ["Prototype.html", "Interactive Prototype", "~38 KB / single file", "5 fully interactive screens (May 23, 2026)"],
    ["Test_Validation_Report.docx", "Test Report v1.0", "10 pages", "32 functional + 8 NFR + 9 security test results (June 13, 2026)"],
    ["Final_Presentation.pptx", "Presentation v1.0", "18 slides", "Live demo + speaker notes (June 13, 2026)"],
    ["Presentation_Script.docx", "Speaker Script", "8 pages", "Full per-slide script + Q&A preparation"],
    ["Final_Project_Report.docx", "This document", "23 pages", "Comprehensive integration of all deliverables"],
]
add_table(doc, deliv_a_headers, deliv_a_rows, [2.4, 1.4, 1.0, 2.5])

body_no_indent(doc, "")
body_no_indent(doc, "All files have been compressed into SAD_Project_Submission.zip and submitted to the course Moodle page.")
body_no_indent(doc, "")
body_no_indent(doc, "The complete project is also publicly available on GitHub:")
body_no_indent(doc, "    https://github.com/a2318491287-design/campus-resource-platform")
body_no_indent(doc, "")
body_no_indent(doc, "The repository serves as our project management evidence (per the BBAZ16604 requirement to use software for managing the group). It contains:")
bullet(doc, "All source code (FastAPI backend + interactive HTML prototype + 8 documentation generation scripts)")
bullet(doc, "All deliverable documents (Charter, Requirements, Design, Reports, Test Report, Final Report)")
bullet(doc, "8 closed Issues representing the phase-by-phase task management workflow (Phase 1 user research through Phase 5 final reporting)")
bullet(doc, "Initial commit log capturing the full project state at submission")
bullet(doc, "README.md with complete project navigation, deployment instructions, and team attribution")

doc.add_page_break()

# ========== APPENDIX B ==========
heading1(doc, "Appendix B — Acronyms and Glossary")

acronyms_headers = ["Term / Acronym", "Definition"]
acronyms_rows = [
    ["RAD", "Requirements Analysis Document"],
    ["SDD", "System Design Document"],
    ["FR", "Functional Requirement"],
    ["NFR", "Non-Functional Requirement"],
    ["UC", "Use Case"],
    ["DFD", "Data Flow Diagram"],
    ["ERD", "Entity-Relationship Diagram"],
    ["RTM", "Requirements Traceability Matrix"],
    ["UML", "Unified Modeling Language"],
    ["SAD", "System Analysis and Design (course context)"],
    ["MoSCoW", "Must Have / Should Have / Could Have / Won't Have prioritization framework"],
    ["MVP", "Minimum Viable Product"],
    ["PIPL", "Personal Information Protection Law (PRC)"],
    ["WCAG", "Web Content Accessibility Guidelines"],
    ["JWT", "JSON Web Token (authentication)"],
    ["CSRF", "Cross-Site Request Forgery (security threat)"],
    ["XSS", "Cross-Site Scripting (security threat)"],
    ["NPS", "Net Promoter Score (user satisfaction metric)"],
    ["TC", "Test Case"],
    ["DD", "Design Decision (logged)"],
    ["KPI", "Key Performance Indicator"],
    ["p95", "95th percentile (latency measurement)"],
    ["FOR UPDATE", "MySQL row-level locking clause used in atomic transactions"],
    ["FULLTEXT INDEX", "MySQL full-text search index supporting natural-language queries"],
    ["Resource", "Any academic file uploaded to the platform — notes, past papers, assignments, other"],
    ["Tag", "Searchable keyword label assigned to a resource"],
    ["Points (积分)", "Virtual currency earned by users for sharing contributions"],
    ["Pin (置顶)", "Premium reward placing a resource at the top of relevant search results for 7 days"],
    ["Free Daily Downloads", "Three downloads per user per day at zero point cost, resetting at midnight"],
]
add_table(doc, acronyms_headers, acronyms_rows, [2.0, 4.5])

doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph("End of Final Project Report  ·  Version 1.0  ·  June 20, 2026  ·  Campus Academic Resource Sharing Platform")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(10)
footer.runs[0].font.bold = True
footer.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

footer2 = doc.add_paragraph("System Analysis and Design  ·  Spring 2026  ·  School of Business  ·  Macau University of Science and Technology")
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer2.runs[0].font.size = Pt(9)
footer2.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save('/Users/yuxianglian/Documents/系统分析与设计/SAD_Project/Final_Project_Report.docx')
print("Done: Final_Project_Report.docx")
