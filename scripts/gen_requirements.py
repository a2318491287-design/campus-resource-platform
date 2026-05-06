from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

def set_style(paragraph, size=12, bold=False, color=None, align=None):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        paragraph.alignment = align

def heading1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(15)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E74B5')
        shd.set(qn('w:color'), 'FFFFFF')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'D6E4F0')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ===================== COVER PAGE =====================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("REQUIREMENTS ANALYSIS DOCUMENT")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Campus Academic Resource Sharing Platform")
r.font.size = Pt(16)
r.font.bold = True

doc.add_paragraph()
info = [
    ("Project Title:", "Campus Academic Resource Sharing Platform Development"),
    ("Document Version:", "1.0"),
    ("Prepared by:", "Lian Yuxiang (1230020693)"),
    ("Team Members:", "Lian Yuxiang 1230020693  |  Yu Kaijie 1230020426  |  Chen Hanzhong 1230032209"),
    ("Course:", "System Analysis and Design"),
    ("Lecturer:", "Dr. CHE Pak Hou (Howard)"),
    ("Date:", "April 4, 2026"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}  ")
    r.font.size = Pt(11)
    r.font.bold = True
    r2 = p.add_run(value)
    r2.font.size = Pt(11)

doc.add_page_break()

# ===================== TABLE OF CONTENTS =====================
heading1(doc, "Table of Contents")
toc_items = [
    "1. Introduction .................................................. 3",
    "2. Business Background & Problem Statement ........................ 3",
    "3. Stakeholder Analysis .......................................... 4",
    "4. Functional Requirements ....................................... 5",
    "   4.1 Feature Improvement: Optimized Precise Retrieval .......... 5",
    "   4.2 New Feature: Points-Based Incentive System ............... 6",
    "5. Non-Functional Requirements .................................. 7",
    "6. Use Case Diagrams & Descriptions ............................. 8",
    "7. Data Flow Diagrams (DFD) .................................... 11",
    "8. Entity-Relationship Diagram (ERD) ........................... 12",
    "9. Requirements Traceability Matrix ............................ 13",
    "10. Assumptions & Constraints .................................. 14",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ===================== SECTION 1: INTRODUCTION =====================
heading1(doc, "1. Introduction")
heading2(doc, "1.1 Purpose of This Document")
body(doc, "This Requirements Analysis Document (RAD) defines the complete set of functional and non-functional requirements for the Campus Academic Resource Sharing Platform. It serves as the authoritative specification for the system design, prototype development, and testing phases, following ISO 29148 requirements engineering standards.")

heading2(doc, "1.2 Project Overview")
body(doc, "The Campus Academic Resource Sharing Platform is a mobile/web application designed to address the fragmented academic resource ecosystem within university campuses. The project targets two core improvements:")
bullet(doc, "Feature Improvement: An optimized precise retrieval module that enables students to locate academic materials (lecture notes, past papers, assignments) with significantly higher efficiency than current search approaches.")
bullet(doc, "New Feature: A points-based incentive system that motivates students to contribute high-quality resources by rewarding sharing behaviors with redeemable points.")

heading2(doc, "1.3 Scope")
body(doc, "This document covers requirements for the following system boundaries:")
bullet(doc, "In Scope: User registration/login, resource upload, multi-dimensional search and filtering, points earning and redemption, resource rating and review, progress tracking dashboard.")
bullet(doc, "Out of Scope: Backend API implementation, mobile app store deployment, iOS 14 and earlier / Android 8.0 and earlier, non-academic content (e.g., food delivery, social forum), off-campus users.")

heading2(doc, "1.4 Definitions and Acronyms")
headers = ["Term", "Definition"]
rows = [
    ["RAD", "Requirements Analysis Document"],
    ["UC", "Use Case"],
    ["DFD", "Data Flow Diagram"],
    ["ERD", "Entity-Relationship Diagram"],
    ["FR", "Functional Requirement"],
    ["NFR", "Non-Functional Requirement"],
    ["Points", "Virtual currency earned by users for contributing resources"],
    ["Resource", "Any academic file (notes, papers, assignments, past exams) uploaded to the platform"],
    ["Tag", "Keyword label assigned to a resource for categorization and search filtering"],
]
add_table(doc, headers, rows, [1.5, 4.5])

# ===================== SECTION 2: BUSINESS BACKGROUND =====================
doc.add_paragraph()
heading1(doc, "2. Business Background & Problem Statement")
heading2(doc, "2.1 Current Situation")
body(doc, "University students currently rely on a fragmented set of channels to obtain academic resources: group chats (WeChat/WhatsApp), personal file drives, paid third-party platforms (e.g., Baidu Wenku), or direct peer requests. This creates the following documented pain points:")

pain_headers = ["Pain Point", "Description", "Impact"]
pain_rows = [
    ["Resource Fragmentation", "Materials scattered across multiple platforms with no unified access point", "High — students waste 30-60 min per search session"],
    ["Low Retrieval Precision", "Generic keyword search returns irrelevant results; no course/year filtering", "High — wrong-year past papers frequently downloaded"],
    ["Hoarding Behavior", "No incentive for sharing; students withhold materials to maintain competitive advantage", "Medium — community resource pool depletes over time"],
    ["Inflated Second-Hand Prices", "Premium resources sold at unfair prices on peer-to-peer channels", "Medium — financially burdens lower-income students"],
    ["No Quality Assurance", "No rating or validation mechanism; inaccurate/outdated materials circulate unchecked", "High — risk of studying incorrect content before exams"],
]
add_table(doc, pain_headers, pain_rows, [1.8, 3.0, 2.2])

heading2(doc, "2.2 Proposed Solution")
body(doc, "The platform consolidates academic resources into a single searchable repository with two differentiating features: (1) a multi-dimensional smart retrieval engine that filters by course code, year, material type, and quality rating; and (2) a gamified points system that converts sharing contributions into tangible benefits (download credits, priority display).")

heading2(doc, "2.3 Business Objectives")
bullet(doc, "Reduce average resource retrieval time by 70% compared to current fragmented channels.")
bullet(doc, "Increase voluntary sharing rate by 80% through points-based incentives within one semester of deployment.")
bullet(doc, "Achieve overall user satisfaction score of 4.0/5.0 or above in post-deployment surveys.")
bullet(doc, "Ensure 95% functional requirement coverage verified through prototype user testing.")

# ===================== SECTION 3: STAKEHOLDER ANALYSIS =====================
doc.add_paragraph()
heading1(doc, "3. Stakeholder Analysis")
body(doc, "The following stakeholders have been identified based on their influence on and interest in the platform, following a Power-Interest Matrix framework.")

s_headers = ["Stakeholder", "Role", "Interest", "Influence", "Key Needs"]
s_rows = [
    ["Dr. CHE Pak Hou", "Course Lecturer / Project Sponsor", "High", "High", "Complete deliverables on time; academic rigor; ISO compliance"],
    ["Course TA", "Evaluation & Feedback", "Medium", "High", "Clear documentation; testable requirements; traceable design"],
    ["Undergraduate Students", "Primary End Users", "High", "Medium", "Fast search; reliable quality; fair reward for contributions"],
    ["University Student Union", "Potential Deployment Partner", "Medium", "Medium", "Platform aligns with academic integrity policies"],
    ["Project Team (3 members)", "Developers / Analysts", "High", "High", "Feasible scope; clear task division; course score"],
    ["Peer Project Teams", "Indirect Benchmarking", "Low", "Low", "—"],
]
add_table(doc, s_headers, s_rows, [1.4, 1.6, 0.8, 0.9, 2.3])

# ===================== SECTION 4: FUNCTIONAL REQUIREMENTS =====================
doc.add_paragraph()
heading1(doc, "4. Functional Requirements")
body(doc, "Requirements are numbered FR-XX and prioritized using MoSCoW: Must Have (M), Should Have (S), Could Have (C), Won't Have (W).")

heading2(doc, "4.1 User Account Management")
uc_headers = ["ID", "Requirement", "Priority", "Rationale"]
uc_rows = [
    ["FR-01", "The system shall allow students to register using their university student ID and email.", "M", "Gate access to verified campus community"],
    ["FR-02", "The system shall authenticate users via student ID and password combination.", "M", "Security and identity assurance"],
    ["FR-03", "The system shall display a user profile showing username, points balance, upload count, and download count.", "M", "Engagement and transparency"],
    ["FR-04", "The system shall allow users to reset their password via registered email.", "S", "Account recovery usability"],
]
add_table(doc, uc_headers, uc_rows, [0.7, 3.5, 0.9, 2.0])

heading2(doc, "4.2 Feature Improvement: Optimized Precise Retrieval Module")
body(doc, "This module improves upon the baseline keyword search by adding structured metadata filtering and relevance ranking.")
r_rows = [
    ["FR-05", "The system shall support keyword search across resource titles, descriptions, and tags.", "M", "Core search functionality"],
    ["FR-06", "The system shall allow filtering by Course Code (e.g., BBAZ16604), Academic Year, Resource Type (Notes / Past Paper / Assignment / Other), and Minimum Rating.", "M", "Key differentiation from generic search"],
    ["FR-07", "The system shall rank results by a composite relevance score: match accuracy (40%) + download count (30%) + average rating (30%).", "M", "Surface highest-quality materials first"],
    ["FR-08", "The system shall display a preview snippet (first 200 characters or thumbnail) in search results without requiring full download.", "S", "Reduce unnecessary downloads"],
    ["FR-09", "The system shall maintain a personal search history for each user (last 20 searches).", "S", "Re-access convenience"],
    ["FR-10", "The system shall recommend related resources based on the currently viewed resource's tags and course code.", "C", "Discovery and engagement"],
]
add_table(doc, uc_headers, r_rows, [0.7, 3.5, 0.9, 2.0])

heading2(doc, "4.3 New Feature: Points-Based Incentive System")
body(doc, "This module introduces a gamified economy to incentivize resource sharing and quality contributions.")
p_rows = [
    ["FR-11", "The system shall award points to users upon successful resource upload: 10 points per upload (pending admin approval).", "M", "Core incentive for sharing"],
    ["FR-12", "The system shall award 2 points each time a user's uploaded resource is downloaded by another user.", "M", "Reward continued contribution"],
    ["FR-13", "The system shall award 1 point each time a user's resource receives a 4-star or above rating.", "M", "Incentivize quality"],
    ["FR-14", "The system shall deduct 5 points from a user's balance per download (unless balance is zero, in which case 0 deduction for first 3 free downloads per day).", "M", "Create balanced economy"],
    ["FR-15", "The system shall allow users to redeem 50 points for 10 additional free downloads.", "M", "Tangible redemption value"],
    ["FR-16", "The system shall allow users to spend 100 points to pin their uploaded resource to the top of relevant search results for 7 days.", "S", "Premium incentive for large contributors"],
    ["FR-17", "The system shall display a campus leaderboard of top 20 contributors by points earned this month.", "S", "Social motivation and recognition"],
    ["FR-18", "The system shall provide a full transaction history of all point earnings and deductions.", "M", "Transparency and trust"],
    ["FR-19", "The system shall send in-app notifications when a user earns points (upload approved, resource downloaded, rating received).", "S", "Real-time engagement feedback"],
]
add_table(doc, uc_headers, p_rows, [0.7, 3.5, 0.9, 2.0])

heading2(doc, "4.4 Resource Management")
rm_rows = [
    ["FR-20", "The system shall allow authenticated users to upload resources in PDF, DOCX, PPTX, or image formats (max 50MB per file).", "M", "Core upload functionality"],
    ["FR-21", "The system shall require uploaders to tag resources with: Course Code, Year, Type, and at least 2 keywords.", "M", "Metadata quality for search"],
    ["FR-22", "The system shall support a simple admin/peer review queue: uploaded resources appear as 'Pending' until verified.", "S", "Quality assurance gate"],
    ["FR-23", "The system shall allow users to rate resources on a 1-5 star scale and leave a text comment.", "M", "Community quality signaling"],
    ["FR-24", "The system shall allow the original uploader to edit metadata (title, tags, description) of their own resources.", "S", "Content maintenance"],
    ["FR-25", "The system shall allow users to report resources for inaccuracy or copyright violation.", "S", "Policy compliance"],
]
add_table(doc, uc_headers, rm_rows, [0.7, 3.5, 0.9, 2.0])

# ===================== SECTION 5: NON-FUNCTIONAL REQUIREMENTS =====================
doc.add_paragraph()
heading1(doc, "5. Non-Functional Requirements")
nf_headers = ["ID", "Category", "Requirement", "Measurement Criterion"]
nf_rows = [
    ["NFR-01", "Performance", "Search results shall be returned within 2 seconds for 95% of queries under normal load (up to 500 concurrent users).", "Load test with JMeter or equivalent"],
    ["NFR-02", "Performance", "File upload of up to 20MB shall complete within 10 seconds on a standard campus WiFi connection (≥10Mbps).", "Upload time measurement"],
    ["NFR-03", "Usability", "A first-time user shall be able to complete a resource search and download within 3 minutes without external guidance.", "Usability test with 5 student participants"],
    ["NFR-04", "Usability", "The interface shall comply with WCAG 2.1 AA accessibility guidelines.", "Automated accessibility checker (axe)"],
    ["NFR-05", "Reliability", "The system shall achieve 99% uptime during semester peak periods (exam weeks).", "Uptime monitoring log"],
    ["NFR-06", "Security", "All user passwords shall be stored using bcrypt hashing with salt (cost factor ≥ 12).", "Code review / security audit"],
    ["NFR-07", "Security", "The system shall comply with PRC Personal Information Protection Law (PIPL) and MUST university data privacy policy.", "Legal review checklist"],
    ["NFR-08", "Security", "All data transmission shall use HTTPS (TLS 1.2 or higher).", "SSL certificate verification"],
    ["NFR-09", "Scalability", "The system architecture shall support horizontal scaling to accommodate up to 5,000 registered users without redesign.", "Architecture review"],
    ["NFR-10", "Maintainability", "All code shall follow documented naming conventions; test coverage shall be ≥ 70% for core modules.", "Code review + coverage report"],
    ["NFR-11", "Compatibility", "The web interface shall support Chrome 90+, Safari 14+, Firefox 88+; mobile support for iOS 15+ and Android 9+.", "Cross-browser testing matrix"],
    ["NFR-12", "Localization", "The primary interface language shall be Simplified Chinese; English labels required for all course codes and file types.", "UI review"],
]
add_table(doc, nf_headers, nf_rows, [0.7, 1.3, 3.2, 1.8])

# ===================== SECTION 6: USE CASES =====================
doc.add_paragraph()
heading1(doc, "6. Use Case Diagrams & Descriptions")

heading2(doc, "6.1 System Use Case Overview")
body(doc, "The following actors interact with the system:")
bullet(doc, "Student User: Registers, logs in, searches resources, downloads, uploads, rates, earns/redeems points.")
bullet(doc, "Guest (Unauthenticated): Can browse resource titles only; cannot download or upload.")
bullet(doc, "Admin Reviewer: Approves/rejects uploaded resources; manages reported content.")
bullet(doc, "System (Automated): Awards points upon triggered events; sends notifications.")

body(doc, "Primary Use Cases:")
bullet(doc, "UC-01: Register Account")
bullet(doc, "UC-02: Login / Logout")
bullet(doc, "UC-03: Search Resources (with multi-dimensional filter)")
bullet(doc, "UC-04: Download Resource")
bullet(doc, "UC-05: Upload Resource")
bullet(doc, "UC-06: Rate & Review Resource")
bullet(doc, "UC-07: View Points Balance & Transaction History")
bullet(doc, "UC-08: Redeem Points")
bullet(doc, "UC-09: View Leaderboard")
bullet(doc, "UC-10: Admin Review Resource Queue")

heading2(doc, "6.2 Use Case Descriptions")

def use_case_block(doc, uc_id, name, actor, precond, trigger, main_flow, alt_flow, postcond, reqs):
    heading3(doc, f"{uc_id}: {name}")
    uc_data = [
        ["Use Case ID", uc_id],
        ["Use Case Name", name],
        ["Primary Actor", actor],
        ["Preconditions", precond],
        ["Trigger", trigger],
        ["Main Success Scenario", main_flow],
        ["Alternative / Exception Flows", alt_flow],
        ["Postconditions", postcond],
        ["Related Requirements", reqs],
    ]
    table = doc.add_table(rows=len(uc_data), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(uc_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[0].width = Inches(1.8)
        row.cells[1].text = value
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].width = Inches(4.7)
    doc.add_paragraph()

use_case_block(doc,
    "UC-03", "Search Resources (Optimized Precise Retrieval)",
    "Student User (authenticated or guest for browse)",
    "User is on the Search page.",
    "User enters a keyword or applies one or more filters.",
    "1. User enters keyword in search bar.\n2. User optionally selects filters: Course Code, Year, Type, Min Rating.\n3. System queries the resource index and computes relevance scores.\n4. System returns ranked results within 2 seconds.\n5. Each result card shows: title, course code, type, year, rating, download count, and preview snippet.\n6. User clicks a result to view full detail page.\n7. User clicks Download (authenticated) → system checks points balance → deducts 5 points → initiates download.",
    "3a. No results found → system displays 'No results' message with suggested alternative keywords.\n7a. User has insufficient points and has exhausted 3 free daily downloads → system displays 'Insufficient Points' dialog with link to upload resources to earn more.\n7b. User is guest → system prompts login/register before allowing download.",
    "Resource file downloaded to user device. Points deducted from user balance. Download count incremented on resource record.",
    "FR-05, FR-06, FR-07, FR-08, FR-09, FR-14, NFR-01"
)

use_case_block(doc,
    "UC-05", "Upload Resource",
    "Student User (authenticated)",
    "User is logged in and on the Upload page.",
    "User clicks the 'Upload Resource' button.",
    "1. User selects file (PDF/DOCX/PPTX/image, max 50MB).\n2. User fills in mandatory metadata: Title, Course Code, Academic Year, Resource Type, and at least 2 keyword tags.\n3. User optionally adds a description (max 500 characters).\n4. User submits the upload form.\n5. System validates file format and size.\n6. System creates a resource record with status = 'Pending'.\n7. System adds resource to the admin review queue.\n8. System notifies user: 'Your resource has been submitted for review.'\n9. Upon admin approval: resource status changes to 'Published'; system awards 10 points to uploader; system sends in-app notification.",
    "5a. File exceeds 50MB → system displays error: 'File too large. Maximum size is 50MB.'\n5b. Invalid file format → system displays error with accepted formats list.\n9a. Admin rejects resource → system notifies uploader with rejection reason; no points awarded.",
    "New resource record created in database. If approved: resource visible in search; 10 points added to uploader's balance.",
    "FR-20, FR-21, FR-22, FR-11, FR-19"
)

use_case_block(doc,
    "UC-08", "Redeem Points",
    "Student User (authenticated)",
    "User has sufficient points balance (≥50 points).",
    "User navigates to the Points Dashboard and clicks 'Redeem'.",
    "1. User views available redemption options: 50 pts → 10 extra downloads; 100 pts → 7-day resource pin.\n2. User selects a redemption option.\n3. System displays confirmation dialog showing points cost and benefit.\n4. User confirms redemption.\n5. System deducts points from user balance.\n6. System applies the reward (increments download credits or activates resource pin).\n7. System records the transaction in user's points history.\n8. System displays success confirmation.",
    "4a. User cancels → no change.\n5a. Between confirmation and deduction, user balance drops below threshold (race condition) → system re-validates; displays 'Insufficient points' and cancels transaction.",
    "User's points balance reduced by redemption cost. Reward activated and visible on user profile.",
    "FR-15, FR-16, FR-18"
)

# ===================== SECTION 7: DFD =====================
doc.add_paragraph()
heading1(doc, "7. Data Flow Diagrams (DFD)")
heading2(doc, "7.1 Context Diagram (Level 0)")
body(doc, "The Level 0 DFD shows the system as a single process interacting with external entities:")
bullet(doc, "External Entities: Student User, Admin Reviewer, Email Notification Service, File Storage Service")
bullet(doc, "Data flows INTO the system: Login credentials, Search queries, Upload files + metadata, Rating submissions, Redemption requests")
bullet(doc, "Data flows OUT OF the system: Search results, Downloaded files, Points balance updates, Notification messages, Approval/rejection status")

body(doc, "[Context Diagram — to be rendered in draw.io]", indent=True)
body(doc, "    ┌──────────────┐                    ┌─────────────────────────────┐", indent=True)
body(doc, "    │ Student User │ ──── queries ──────▶│                             │", indent=True)
body(doc, "    │              │ ◀─── results ───────│   Campus Resource Platform  │", indent=True)
body(doc, "    │              │ ──── uploads ───────▶│        (Main System)        │", indent=True)
body(doc, "    └──────────────┘                    │                             │", indent=True)
body(doc, "    ┌──────────────┐                    │                             │", indent=True)
body(doc, "    │    Admin     │ ──── approvals ────▶│                             │", indent=True)
body(doc, "    └──────────────┘                    └─────────────────────────────┘", indent=True)

heading2(doc, "7.2 Level 1 DFD — Resource Management Subsystem")
body(doc, "Processes decomposed from the main system:")

l1_headers = ["Process No.", "Process Name", "Input Data Flows", "Output Data Flows", "Data Stores Used"]
l1_rows = [
    ["P1", "User Authentication", "Student ID, Password", "Auth token, Error message", "D1: Users"],
    ["P2", "Resource Search & Retrieval", "Search keyword, Filter parameters", "Ranked result list, Preview snippets", "D2: Resources, D3: Tags"],
    ["P3", "Resource Upload & Validation", "File, Metadata form", "Upload confirmation, Review queue entry", "D2: Resources, D3: Tags"],
    ["P4", "Admin Review", "Review decision (approve/reject)", "Status update, Notification trigger", "D2: Resources"],
    ["P5", "Points Calculation & Award", "Trigger events (download, rating, upload approval)", "Points transaction record", "D4: PointRecords, D1: Users"],
    ["P6", "Points Redemption", "Redemption request, Points amount", "Reward activation, Transaction record", "D4: PointRecords, D1: Users"],
    ["P7", "Rating & Review", "Star rating, Comment text, Resource ID", "Updated resource rating, Points award trigger", "D2: Resources, D4: PointRecords"],
]
add_table(doc, l1_headers, l1_rows, [0.8, 1.4, 1.6, 1.6, 1.6])

body(doc, "\nData Stores:")
bullet(doc, "D1: Users — Stores user account information, authentication data, and points balance")
bullet(doc, "D2: Resources — Stores all resource metadata, status, ratings, and download counts")
bullet(doc, "D3: Tags — Stores tag definitions and resource-tag relationships")
bullet(doc, "D4: PointRecords — Stores all point transaction history (earn/spend events)")

# ===================== SECTION 8: ERD =====================
doc.add_paragraph()
heading1(doc, "8. Entity-Relationship Diagram (ERD)")
heading2(doc, "8.1 Entity Definitions")
body(doc, "The following entities form the core data model of the platform:")

erd_headers = ["Entity", "Attributes", "Primary Key"]
erd_rows = [
    ["User", "user_id, student_id, username, password_hash, email, points_balance, upload_count, download_credits, created_at, is_admin", "user_id"],
    ["Resource", "resource_id, title, description, file_path, file_type, file_size, course_code, academic_year, resource_type, status (pending/published/rejected/removed), avg_rating, download_count, uploader_id, created_at, updated_at", "resource_id"],
    ["Tag", "tag_id, tag_name, category (course/type/keyword)", "tag_id"],
    ["ResourceTag", "resource_id (FK), tag_id (FK)", "Composite (resource_id, tag_id)"],
    ["Rating", "rating_id, resource_id (FK), user_id (FK), stars (1-5), comment, created_at", "rating_id"],
    ["PointRecord", "record_id, user_id (FK), resource_id (FK, nullable), action_type (UPLOAD/DOWNLOAD_RECEIVED/RATING_RECEIVED/SPEND_DOWNLOAD/REDEEM), points_delta, balance_after, created_at", "record_id"],
    ["Download", "download_id, resource_id (FK), user_id (FK), downloaded_at", "download_id"],
    ["Redemption", "redemption_id, user_id (FK), reward_type (DOWNLOAD_CREDIT/PIN), points_cost, activated_at, expires_at", "redemption_id"],
]
add_table(doc, erd_headers, erd_rows, [1.3, 4.0, 1.7])

heading2(doc, "8.2 Relationships")
rel_headers = ["Relationship", "Entities Involved", "Cardinality", "Description"]
rel_rows = [
    ["uploads", "User → Resource", "1 : M", "One user can upload many resources; each resource has one uploader"],
    ["has_tags", "Resource ↔ Tag", "M : M (via ResourceTag)", "A resource can have multiple tags; a tag can apply to multiple resources"],
    ["rates", "User → Resource (via Rating)", "M : M", "A user can rate many resources; a resource can receive ratings from many users (one rating per user per resource)"],
    ["downloads", "User → Resource (via Download)", "M : M", "A user can download many resources; a resource can be downloaded by many users"],
    ["earns/spends", "User → PointRecord", "1 : M", "All point changes for a user are recorded as individual transactions"],
    ["redeems", "User → Redemption", "1 : M", "A user can make multiple redemptions over time"],
]
add_table(doc, rel_headers, rel_rows, [1.2, 1.8, 1.5, 2.5])

# ===================== SECTION 9: RTM =====================
doc.add_paragraph()
heading1(doc, "9. Requirements Traceability Matrix (RTM)")
body(doc, "This matrix traces each functional requirement to its originating stakeholder need and the corresponding system design component (to be updated as design progresses).")

rtm_headers = ["Req. ID", "Requirement Summary", "Stakeholder Need", "Use Case", "Design Component", "Test Case"]
rtm_rows = [
    ["FR-05", "Keyword search", "Students need fast resource discovery", "UC-03", "Search Engine Module", "TC-01"],
    ["FR-06", "Multi-dimensional filter", "Students waste time on irrelevant results", "UC-03", "Filter Component", "TC-02"],
    ["FR-07", "Relevance ranking", "Surface high-quality materials first", "UC-03", "Ranking Algorithm", "TC-03"],
    ["FR-11", "Award 10 pts on upload", "No incentive to share resources", "UC-05", "Points Engine", "TC-04"],
    ["FR-12", "Award 2 pts per download", "Reward continued contribution", "UC-04", "Points Engine", "TC-05"],
    ["FR-14", "Deduct 5 pts per download", "Balanced economy sustainability", "UC-04", "Points Engine", "TC-06"],
    ["FR-15", "Redeem 50 pts → 10 downloads", "Tangible redemption value", "UC-08", "Redemption Module", "TC-07"],
    ["FR-17", "Monthly leaderboard", "Social motivation", "UC-09", "Leaderboard Component", "TC-08"],
    ["FR-20", "File upload (PDF/DOCX/PPTX)", "Core upload capability", "UC-05", "File Upload Module", "TC-09"],
    ["FR-23", "1-5 star rating + comment", "Community quality signaling", "UC-06", "Rating Module", "TC-10"],
]
add_table(doc, rtm_headers, rtm_rows, [0.7, 1.8, 1.6, 0.7, 1.5, 0.7])

# ===================== SECTION 10: ASSUMPTIONS & CONSTRAINTS =====================
doc.add_paragraph()
heading1(doc, "10. Assumptions & Constraints")
heading2(doc, "10.1 Assumptions")
bullet(doc, "All target users possess a valid university student ID and a campus email address.")
bullet(doc, "Campus WiFi infrastructure supports concurrent file transfers at ≥10Mbps per active user session.")
bullet(doc, "Students are willing to participate in usability testing sessions (3-5 volunteers available).")
bullet(doc, "The university does not prohibit the operation of a peer-to-peer academic resource platform.")
bullet(doc, "Resources shared on the platform are copyright-compliant (course materials with lecturer permission, student-created notes).")
bullet(doc, "The project team has continuous access to Figma, draw.io, Git/GitHub, and MySQL during the project duration.")

heading2(doc, "10.2 Constraints")
bullet(doc, "Timeline: Requirements phase completed by April 4, 2026; full project by June 30, 2026.")
bullet(doc, "Team size: Fixed at 3 members; no additional resources available.")
bullet(doc, "Budget: Total $10,000 HKD; no paid cloud infrastructure; free-tier tools only.")
bullet(doc, "Technology: Must use MySQL for data storage and Git/GitHub for version control (per course requirements).")
bullet(doc, "Regulatory: Must comply with PRC Personal Information Protection Law (PIPL) and university data retention policies.")
bullet(doc, "Scope freeze: Requirements may not be modified after Week 5 without written approval from Dr. CHE.")

# ===================== FOOTER =====================
doc.add_paragraph()
doc.add_paragraph()
footer_p = doc.add_paragraph("End of Requirements Analysis Document  |  Version 1.0  |  April 4, 2026  |  System Analysis and Design — MUST")
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.runs[0].font.size = Pt(9)
footer_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

import os
os.makedirs('/Users/yuxianglian/Downloads/SAD_Project', exist_ok=True)
doc.save('/Users/yuxianglian/Downloads/SAD_Project/Requirements_Analysis_Document.docx')
print("Done: Requirements_Analysis_Document.docx")
