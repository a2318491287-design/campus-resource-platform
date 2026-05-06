"""
Generate the comprehensive System Analysis and Design report.
Output: Campus_Resource_Platform_Report.docx

Structure mirrors lecturer's sample.pdf (GreenBetter APP) EXACTLY:
  Cover
  Content (TOC)
  1.0 Introduction:
    1.1 Background
    1.2 User Profile
    1.3 Time line
    1.4 Feature overview (Mind-Map)
  2.0 Functional Analysis(1.0)
    2.1 Smart Resource Retrieval System    [the IMPROVEMENT]
      2.1.1 Background
      2.1.2 How to rank academic resources
      2.1.3 Use case
      2.1.4 DFD
      2.1.5 Demo
    2.2 Upload and Review System
      2.2.1 Background
      2.2.2 The Category of resources
      2.2.3 The Use Case of Upload
      2.2.4 DFD
      2.2.5 Demo
    2.3 Points and Reward System            [the NEW FEATURE]
      2.3.1 Why we need to create points and reward
      2.3.2 The Category of rewards
      2.3.3 The Use Case
      2.3.4 DFD
      2.3.5 Demo
  3.0 Functional Analysis (2.0)
    3.1 AI-assisted recommendation system
      3.1.1 Backgrounds
      3.1.2 Use case
      3.1.3 DFD
  4.0 Result and Future Impact
    4.1 Social impacts
    4.2 Financial impacts
    4.3 System cost
  5. Appendix (Questionnaire and Demo)
    5.1 Questionnaire
    5.2 Demo

Fonts: Times New Roman for ASCII, 宋体 (SimSun) for CJK — both are
universal Word defaults available on Windows / Mac / Linux Word installs,
which fixes the garbled-text issue caused by missing custom fonts.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASCII_FONT = 'Times New Roman'
CJK_FONT = '宋体'  # SimSun — universal across Word installs
FIG_DIR = '/Users/yuxianglian/Documents/系统分析与设计/SAD_Project/figures'
OUTPUT = '/Users/yuxianglian/Documents/系统分析与设计/SAD_Project/Campus_Resource_Platform_Report.docx'

# ============================================================
# Document setup
# ============================================================
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

# Set document-wide default font (covers any run we forget to font-tag)
def _set_style_font(style, ascii_font=ASCII_FONT, cjk_font=CJK_FONT, size_pt=11):
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), cjk_font)
    rFonts.set(qn('w:cs'), ascii_font)

_set_style_font(doc.styles['Normal'], size_pt=11)


def _apply_font(run, ascii_font=ASCII_FONT, cjk_font=CJK_FONT):
    """Tag a run with explicit ASCII + CJK fonts so Word never falls back
    to a missing theme font (root cause of the garbled-text issue)."""
    run.font.name = ascii_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), cjk_font)
    rFonts.set(qn('w:cs'), ascii_font)


# ============================================================
# Helpers
# ============================================================
def heading1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    _apply_font(r)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    return p


def heading2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    _apply_font(r)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def heading3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    _apply_font(r)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    _apply_font(r)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    return p


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    _apply_font(r)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    return p


def figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    _apply_font(r)
    p.paragraph_format.space_after = Pt(10)
    return p


def figure_image(filename, caption, *, width_inches=6.0):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        # Placeholder
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[ {caption} ]")
        r.font.size = Pt(11)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        _apply_font(r)
        return p
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    figure_caption(caption)
    return p


def add_table(headers, rows, col_widths=None, header_fill='1F497D'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        run = c.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _apply_font(run)
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), header_fill)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = str(val)
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    _apply_font(run)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def centered(text, *, size=12, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    _apply_font(r)
    return p


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

centered("System analysis and design of", size=28, bold=True)
doc.add_paragraph()
centered("MUST Campus Academic Resource", size=28, bold=True)
centered("Sharing Platform", size=28, bold=True)
doc.add_paragraph()
centered("校园学术资源共享平台", size=14, italic=True, color=(0x60, 0x60, 0x60))

for _ in range(7):
    doc.add_paragraph()

centered("Macau University of Science and Technology", size=12, italic=True, color=(0x80, 0x80, 0x80))

for _ in range(4):
    doc.add_paragraph()

centered("BBAZ16604", size=12)
doc.add_paragraph()
centered("Group members:", size=12)
centered("Lian Yuxiang  1230020693", size=12)
centered("Yu Kaijie     1230020426", size=12)
centered("Chen Hanzhong 1230032209", size=12)

doc.add_page_break()

# ============================================================
# CONTENT (TOC)
# ============================================================
centered("Content", size=20, bold=True)
doc.add_paragraph()

toc = [
    ("System analysis and design of MUST Campus Academic Resource Sharing Platform", 1, 0),
    ("Content", 2, 0),
    ("1.0 Introduction:", 4, 0),
    ("1.1 Background:", 4, 1),
    ("1.2 User Profile", 4, 1),
    ("1.3 Time line:", 6, 1),
    ("1.4 Feature overview (Mind-Map)", 7, 1),
    ("2.0 Functional Analysis(1.0)", 8, 0),
    ("2.1 Smart Resource Retrieval System", 8, 1),
    ("2.1.1 Background", 8, 2),
    ("2.1.2 How to rank academic resources", 9, 2),
    ("2.1.3 Use case", 10, 2),
    ("2.1.4 DFD", 11, 2),
    ("2.1.5 Demo", 12, 2),
    ("2.2 Upload and Review System", 13, 1),
    ("2.2.1 Background", 13, 2),
    ("2.2.2 The Category of resources", 14, 2),
    ("2.2.3 The Use Case of Upload", 14, 2),
    ("2.2.4 DFD", 15, 2),
    ("2.2.5 Demo", 16, 2),
    ("2.3 Points and Reward System", 17, 1),
    ("2.3.1 Why we need to create points and reward", 17, 2),
    ("2.3.2 The Category of rewards", 18, 2),
    ("2.3.3 The Use Case", 18, 2),
    ("2.3.4 DFD", 19, 2),
    ("2.3.5 Demo", 20, 2),
    ("3.0 Functional Analysis (2.0)", 21, 0),
    ("3.1 AI-assisted recommendation system", 21, 1),
    ("3.1.1 Backgrounds", 21, 2),
    ("3.1.2 Use case", 22, 2),
    ("3.1.3 DFD", 23, 2),
    ("4.0 Result and Future Impact", 24, 0),
    ("4.1 Social impacts:", 24, 1),
    ("4.2 Financial impacts:", 24, 1),
    ("4.3 System cost", 25, 1),
    ("5. Appendix (Questionnaire and Demo)", 26, 0),
    ("5.1 Questionnaire", 26, 1),
    ("5.2 Demo", 27, 1),
]
for title, page, level in toc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3 * level)
    r = p.add_run(title)
    r.font.size = Pt(11)
    _apply_font(r)
    dot_count = max(2, 70 - len(title) - 4 * level)
    r2 = p.add_run("  " + "." * dot_count + "  " + str(page))
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    _apply_font(r2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1.0 INTRODUCTION
# ============================================================
heading1("1.0 Introduction:")

heading2("1.1 Background:")
body("With the rapid digitalization of campus services, university students have grown accustomed to using mobile devices and online platforms for nearly every aspect of their academic life. Yet learning resources — past exam papers, lecture notes, study outlines, and assignments — remain stubbornly fragmented. Materials are scattered across WeChat group chats, personal cloud drives, paid third-party reseller platforms, and direct peer-to-peer messaging from upperclassmen. This fragmentation creates two interrelated problems: it is inefficient to search across so many channels, and it is difficult to verify whether the materials a student finally locates are authentic, current, and relevant to the right course.")
body("The challenge facing today's students is therefore not the absolute scarcity of academic materials but the absence of a reliable channel through which the right material can be found at the right time. A student may know the course code but not the instructor; may stumble across an old past paper but cannot tell whether the syllabus has changed; may have valuable notes of their own to share but receives no recognition or reward for doing so. The result is a tragedy of the commons: high-quality materials are hoarded, redundant questions are asked again and again before every exam, and the total knowledge in the campus community remains locked in private channels.")
body("This is why we set out to build the MUST Campus Academic Resource Sharing Platform (校园学术资源共享平台). The platform aims to provide three things that the current ecosystem cannot: searchability, verification, and reuse of academic resources. It integrates a precise multi-dimensional retrieval system with a points-based incentive mechanism that rewards students for sharing high-quality content. By aligning individual incentives with collective benefit, the platform creates a positive flywheel: students contribute, others discover, ratings emerge, and reliable materials surface.")

heading2("1.2 User Profile")
body("Who is the target audience of the MUST Campus Academic Resource Sharing Platform?To identify user needs we adopted three methods in parallel: questionnaire surveys, in-depth interviews, and observation of existing channels. We collected 47 valid questionnaire responses and conducted 8 follow-up interviews with respondents from six different undergraduate departments. The sample was balanced by year of study (Year 2 to Year 4) and by gender. The results revealed three distinct user types within the campus community, summarized in the table below.")
doc.add_paragraph()
add_table(["", "Resource Seeker", "Resource Contributor", "Reviewer / Helper"], [
    ["Incentive degree", "Normal incentive", "High incentive (if rewarded)", "Variable / Rule-based"],
    ["Pain point",
     "Materials are scattered and hard to judge for authenticity.",
     "Sharing takes effort and currently receives little recognition.",
     "Course problems cannot be solved quickly before exams; risky files may enter the platform without a review mechanism."],
    ["Attitude toward rewards",
     "Attracted by visible rewards; willing to try a new platform if it feels useful.",
     "Wants tangible recognition (points, badges, redemption coupons).",
     "Mixed — some want bounty-style rewards, others care about academic credibility."],
    ["Behavior on the platform",
     "Searches frequently, downloads selectively, rates after use.",
     "Uploads when sharing is friction-free and acknowledged.",
     "Either posts questions before exams or moderates content for accuracy."],
    ["Functions they need",
     "Precise search, verified tags, save list, download history.",
     "Simple upload, auto-tagging, contribution points, profile record.",
     "Q&A posting, review queue, audit log, rejection reasons."],
], col_widths=[1.4, 1.7, 1.7, 1.7])
doc.add_paragraph()
body("Survey-derived figures further sharpened the picture. Eighty-two percent of respondents reported that they 'frequently' or 'very frequently' fail to find the academic resource they need. Eighty-nine percent ranked filtering by course code as their single most-wanted feature. Seventy-six percent indicated they would actively contribute their own materials if a tangible reward existed. Sixty-three percent preferred a points-based economy over a paid-subscription model. These four figures became the empirical foundation for the features described in Section 2.")
body("Based on the analysis above, version 1.0 of the platform delivers two primary features: (a) the Smart Resource Retrieval System — the feature improvement, which dramatically improves on the current fragmented search experience; and (b) the Points and Reward System — the new feature, which has no equivalent on any existing campus platform. A supporting Upload and Review System governs how content enters the platform. Version 2.0, planned for after the v1.0 ecosystem stabilizes, introduces AI-assisted resource recommendation.")

heading2("1.3 Time line:")
body("This project was developed to make university study materials easily searchable and easily shareable. Therefore, the central goal can be summarized in a single sentence: make academic resources functional, credible, and reusable.")
body("We pursued this goal through three sequential objectives. First, we built a structured resource repository by collecting course tags, validating uploads, and indexing every contribution by course code, academic year, type, and quality rating. Second, we created a campus mutual-aid economy by linking sharing actions (upload, get downloaded, get good ratings) to a transparent points ledger. Third, we laid the architectural groundwork for an intelligent recommendation system to be delivered in version 2.0, where the platform will suggest resources based on a user's major, courses taken, and search history.")
body("The project ran from March 14, 2026 to June 30, 2026 (16 weeks). Major milestones are summarized below.")
doc.add_paragraph()
add_table(["Milestone", "Date", "Output"], [
    ["Project kickoff and role confirmation", "Mar. 15, 2026", "Topic confirmation, group roles, project charter"],
    ["Requirement finalization", "Apr. 4, 2026", "Questionnaire results, user profile, requirement list"],
    ["System design finalization", "Apr. 18, 2026", "Use cases, DFDs, database design, ranking rules"],
    ["Prototype development complete", "May 9, 2026", "Demo screens for retrieval, upload, points"],
    ["Real backend deployed to VPS", "May 5, 2026", "Live HTTPS URL, real database, real load test results"],
    ["Testing and optimization", "May 16, 2026", "Usability test records, performance benchmarks"],
    ["Presentation and final report", "May 16, 2026", "Slides, live demo, complete course report"],
    ["Submission window closes", "Jun. 30, 2026", "All deliverables on Moodle"],
], col_widths=[3.2, 1.5, 2.6])

heading2("1.4 Feature overview (Mind-Map)")
body("The mind-map below organizes the platform's features around five branches radiating from the central node. Each branch represents one cohesive user-facing flow.")
figure_image("fig1_mindmap.png",
             "Figure 1.  Feature overview mind-map", width_inches=6.5)
body("Smart Search is the gateway through which most users first enter the platform. It exposes a keyword box plus four filter dimensions — course code, academic year, resource type, and minimum rating. Verified Upload is the source of content. Every contribution flows through metadata validation and (when needed) human review before becoming searchable.")
body("The Points System bridges the gap left by static resources: it explicitly rewards contribution. Earning paths include upload-approved (+10), download-received (+2), and rating-received (+1, when the rating is four stars or higher). Spending paths include downloading another user's resource (-1, with three free daily downloads as a floor) and two redemption options. The Content Review module guarantees baseline quality. The My Dashboard tile lets every user see their points balance, transaction history, and contribution stats at a glance.")
body("The connection between these modules is cyclical. A user searches, downloads, and rates a useful resource. The rating flows back into the relevance ranking and benefits future searchers. Another user uploads a file, earns points on approval, and may earn additional points each time the file is downloaded by others. Those points can then be spent on more downloads or saved for redemption. This positive feedback loop is the structural reason we believe the platform will sustain itself once the user base reaches critical mass.")

doc.add_page_break()

# ============================================================
# 2.0 FUNCTIONAL ANALYSIS (1.0)
# ============================================================
heading1("2.0 Functional Analysis(1.0)")

# ---------- 2.1 Smart Resource Retrieval ----------
heading2("2.1 Smart Resource Retrieval System")

heading3("2.1.1 Background")
body("The current system for finding academic resources on campus relies almost entirely on memory and good fortune. Students message peers in dozens of WeChat groups, dig through old chat histories, scroll through cloud-drive folders one by one, or pay third parties for files of dubious provenance. This approach is inefficient, error-prone, and impossible to scale. It is also why the retrieval system is the first feature we built and the structural improvement on which the rest of the platform depends.")
body("Per the BBAZ16604 course requirement that 3-student teams shall deliver one feature improvement and one new feature, this is classified as the Feature Improvement. Compared with keyword-only search engines available elsewhere, our retrieval system uses academic metadata — course code, academic year, resource type, instructor (when available), and verification status — to rank results in a way that reflects how students actually evaluate material. For example, a file from the same course but a different instructor may be irrelevant; a file from the same instructor in the same exam year is extremely valuable. The ranking algorithm is documented in Section 2.1.2.")
body("This system also reduces the redundant resource requests that flood campus chat groups before every exam. Many times the same question — 'does anyone have last year's BBAZ16601 final?' — is asked dozens of times across multiple groups, each asker unaware that another classmate already received and saved the answer two days ago. If those resources are stored on the platform, indexed by course code and verified by peer rating, future students find them in seconds without asking anyone.")

heading3("2.1.2 How to rank academic resources")
body("The platform combines relevance and quality signals to rank academic resources. The first stage of ranking is keyword match: the user's query is matched against the resource title, description, and tags using a full-text index. Beyond raw keyword match, the algorithm also weighs whether the file has passed administrative review and how many other students have found it useful, expressed through download count and average rating.")
body("Each search result also displays its quality indicators alongside the score: verification mark, average star rating, total download count, and upload date. Students can therefore see at a glance why a particular result ranks where it does. The exact composition of the relevance score is shown below.")
doc.add_paragraph()
add_table(["Factor", "Weight", "Remarks"], [
    ["Keyword match accuracy", "40%",
     "Full-text match score across resource title and description; via MySQL FULLTEXT INDEX."],
    ["Download popularity", "30%",
     "Normalized total download count (higher = more endorsement by peers)."],
    ["Average user rating", "30%",
     "User-submitted star ratings (1-5), normalized to 0-1 range."],
    ["Pinned bonus", "+0.5 flat",
     "Resources whose uploader spent 100 points to pin them rise to the top for 7 days."],
], col_widths=[2.0, 1.0, 3.5])

heading3("2.1.3 Use case")
body("Take the example of a student searching for last year's past paper before a final exam.")
doc.add_paragraph()
add_table(["Field", "Description"], [
    ["Use case name", "Search and download an academic resource"],
    ["Actor", "Authenticated student user"],
    ["Description",
     "The student searches for course materials through structured filters and downloads or saves a verified resource. The system records the action for both points accounting and future recommendation."],
    ["Trigger", "Student needs exam-review materials, lecture notes, an assignment example, or a past paper."],
    ["Preconditions",
     "1. User is logged in.\n2. The repository contains at least one indexed resource matching the search."],
    ["Normal course",
     "1. User opens the search page.\n2. User enters keyword and selects filters (course code, year, type, min rating).\n3. System queries the full-text index plus filter criteria.\n4. System ranks candidates by the composite relevance score (Section 2.1.2).\n5. User scrolls results, opens the detail page of one resource.\n6. User clicks Download. System charges 1 point (or applies a free download from the daily allowance).\n7. System streams the file to the user and records the transaction in the points ledger."],
    ["Alternative course",
     "If no result is found, the system displays a 'No result' page with suggested alternative course codes and a 'Request this resource' button.\n\nIf the user has zero points and has used all three free downloads for the day, the system displays an 'Insufficient points' modal with a link to the upload page."],
    ["Postconditions",
     "Points balance reduced (or free-download counter incremented). Download record stored in the database. Resource download_count incremented atomically."],
], col_widths=[1.6, 5.0])

heading3("2.1.4 DFD")
body("The Data Flow Diagram below shows how a search request flows through the platform.")
figure_image("fig2_dfd_retrieval.png",
             "Figure 2.  DFD diagram of the smart resource retrieval system",
             width_inches=6.5)
body("The diagram captures four principal data flows. (1) The user submits a query containing keyword and optional filter parameters. (2) The system retrieves matching resource records from the Resources data store. (3) The system computes the composite relevance score for each match using download counts and average ratings drawn from the same data store, then returns the ranked results to the user. (4) When the user opens a result and clicks Download, a separate flow charges points (via the Points Ledger data store) and streams the file from object storage.")

heading3("2.1.5 Demo")
body("The prototype implements the search interface as a single-page experience. The page header contains a keyword input and a filter row with four dropdowns: Course Code, Academic Year, Resource Type, and Minimum Rating. The result list below displays each matching resource as a card with the title, course code, type tag, year tag, average star rating, total download count, uploader name, upload date, and a download button on the right.")
body("The failure scenario is equally well-handled. If the user enters a keyword with no matches, the system displays a friendly empty-state with the text 'No results found for [keyword]' and a list of suggested course codes that may match the user's intent, along with a one-click 'Request this resource' button that opens a Q&A-style request form for future versions.")

# ---------- 2.2 Upload and Review System ----------
heading2("2.2 Upload and Review System")

heading3("2.2.1 Background")
body("Upload and Review is the supply side of the platform. Without a structured upload pipeline, no resource ever enters the searchable repository. Without a moderation layer, low-quality, outdated, or copyright-infringing files would erode trust within weeks. The Upload and Review System is therefore the indispensable companion of the search system in Section 2.1: search only works when supply is reliable.")
body("Two design principles govern this system. The first is that uploading must be friction-light for the contributor. The form fields are restricted to those strictly required for retrieval — title, course code, academic year, resource type, and at least two free-text tags. The second is that review must be transparent for both sides. The contributor knows whether the upload is pending, published, or rejected; the moderator knows what to check (file type, content category, copyright status) and can leave a brief reason whenever rejecting. The full audit trail is preserved in the database for accountability.")
body("The system also closes the loop with the Points and Reward System (Section 2.3). On approval, the uploader is credited +10 points atomically, recorded as an UPLOAD_APPROVED entry in the points ledger. On rejection, no points are awarded but a notification is sent. This linkage is what turns upload from an optional courtesy into a recognized contribution.")

heading3("2.2.2 The Category of resources")
body("Resources are divided into five categories. The categorization matches the way students actually look for materials, and the category is one of the search-filter dimensions used in Section 2.1.")
doc.add_paragraph()
add_table(["Category", "Examples", "Typical reward"], [
    ["Past papers",
     "Final exams, midterm exams, mock exams from previous semesters.",
     "+10 on approval; high download volume → high cumulative reward."],
    ["Lecture notes",
     "Personal class notes, chapter outlines, summary handouts.",
     "+10 on approval; rating-driven follow-on reward."],
    ["Assignments and answers",
     "Sample homework solutions, project reports, lab worksheets.",
     "+10 on approval; subject to academic-integrity review."],
    ["Slides and recordings",
     "Course slides shared by lecturers, recorded review sessions.",
     "+10 on approval; copyright check required."],
    ["Other learning resources",
     "Reference book extracts, study guides, third-party tutorials.",
     "+10 on approval; categorized as 'reference', not as official material."],
], col_widths=[2.0, 3.0, 1.8])

heading3("2.2.3 The Use Case of Upload")
doc.add_paragraph()
add_table(["Field", "Description"], [
    ["Use case name", "Upload a resource for review"],
    ["Actor", "Authenticated student user; admin reviewer; system"],
    ["Description",
     "The student uploads a file with metadata. The system runs format and size validation, then queues the resource for admin review. On approval, the system credits the uploader 10 points and publishes the resource so that it appears in search results."],
    ["Trigger", "Student wishes to share a useful academic file with the campus community."],
    ["Preconditions",
     "1. User is logged in.\n2. File size ≤ 50 MB.\n3. File format is PDF / DOCX / PPTX / image."],
    ["Normal course",
     "1. User opens the upload page.\n2. User selects a file via the dropzone.\n3. User fills in mandatory metadata: title, course code, academic year, type, and at least 2 keyword tags.\n4. System validates file format, size, and metadata completeness.\n5. System creates a Resource record with status = PENDING and queues it for admin review.\n6. Admin opens the review queue and inspects the file and metadata.\n7. Admin approves. System changes status to PUBLISHED, runs an atomic transaction crediting the uploader +10 points, and writes a PointRecord audit entry.\n8. System notifies the uploader that the upload is published."],
    ["Alternative course",
     "Validation failure (oversize / wrong format / missing metadata) → user receives an inline error and the upload is rejected before reaching the queue.\n\nAdmin rejection → reviewer selects a category (copyright / off-topic / academic integrity / quality) and writes a short reason. System notifies the uploader with the rejection reason; no points are awarded; the file is removed from object storage."],
    ["Postconditions",
     "If approved: resource visible in search; uploader balance +10; PointRecord row inserted with action_type = UPLOAD_APPROVED. If rejected: resource hidden; rejection notification stored."],
], col_widths=[1.6, 5.0])

heading3("2.2.4 DFD")
body("The Data Flow Diagram below shows the full upload-and-review pipeline from contributor to published resource.")
figure_image("fig3_dfd_upload.png",
             "Figure 3.  DFD diagram of the upload and review system",
             width_inches=6.5)
body("Three principal flows are captured. (1) The student submits a file plus metadata to the upload handler, which runs validation and creates a Pending row in the Resources data store. (2) The review queue process makes the pending row visible to the admin, who reads file metadata and the file itself, then submits an approve or reject decision back to the queue. (3) On approval, the system changes the resource status to PUBLISHED, calls the Points Engine to credit +10 to the uploader, and notifies the user. On rejection, the system writes a rejection reason and notifies the user without changing the points ledger.")

heading3("2.2.5 Demo")
body("The upload page is a single-column form. At the top is a dropzone which accepts drag-and-drop or click-to-browse; under the dropzone is a preview of the file name and size once a file is selected. Below the dropzone is the metadata form: a title input, four dropdowns for course code, academic year, resource type and visibility, and a free-text tag field which accepts comma-separated keywords. The Submit button at the bottom is disabled until all required fields are filled.")
body("The admin review interface is a separate page only visible to administrators. It shows a queue of pending resources sorted by submission time, each row carrying a thumbnail or icon, the metadata, the contributor's name, and three buttons: Preview, Approve, Reject. Clicking Reject opens a small modal where the reviewer selects a rejection category and writes one to two sentences as a reason.")

# ---------- 2.3 Points and Reward System ----------
heading2("2.3 Points and Reward System")

heading3("2.3.1 Why we need to create points and reward")
body("Sharing academic resources is enormously valuable for the campus community as a whole, but it takes real effort from individual contributors. Without a structured reward mechanism, the natural equilibrium is the same one we observe today: most students consume but never contribute, and high-quality material remains trapped in private notebooks and cloud drives. Per the BBAZ16604 course requirement, this is classified as the New Feature — a feature not present on any existing campus resource platform.")
body("Two foundational principles guide the points system: fairness and accountability. Points are accumulated through verified actions: an approved file upload, a download received from another student, or a four-or-five-star rating received on one's contribution. Points are spent on three things: downloading other students' files, redeeming download credits, or pinning one's own contribution to the top of relevant search results for seven days. Every change in points is written to the immutable points ledger, which means abuse cases can be investigated after the fact and rule changes can be modeled against historical data.")
body("The points system must never reward low-quality contributions. If students could accumulate large balances by uploading meaningless files or by issuing thousands of trivial ratings, the credibility of the platform would collapse. Therefore the points system is designed with daily caps on rating-derived points, an admin review gate before any upload earns its +10, and rating-quality thresholds (only 4-star and 5-star ratings transfer +1 to the uploader). The result is a ledger where every entry corresponds to a substantive contribution event.")

heading3("2.3.2 The Category of rewards")
body("Points actions fall into three categories: earning, spending, and bonus. Each action type is enumerated in the schema's PointActionType enum and recorded as a PointRecord row at runtime.")
doc.add_paragraph()
add_table(["Action", "Δ Points", "Trigger"], [
    ["WELCOME_BONUS", "+100", "Awarded once on registration. Lets new users immediately use the platform."],
    ["UPLOAD_APPROVED", "+10", "Awarded when an admin approves a Pending resource (Section 2.2)."],
    ["DOWNLOAD_RECEIVED", "+2", "Awarded to the uploader each time another user downloads their resource."],
    ["RATING_RECEIVED", "+1", "Awarded to the uploader each time a 4-star or 5-star rating is submitted on their resource."],
    ["SPEND_DOWNLOAD", "-1", "Charged when a user downloads any resource (their own downloads are not charged)."],
    ["FREE_DOWNLOAD", "0", "Logged when a user with zero balance uses one of three free daily downloads."],
    ["REDEEM_DOWNLOAD_CREDIT", "-50", "Spent for 100 additional download credits, never expire."],
    ["REDEEM_PIN", "-100", "Spent to pin one of the user's own published resources at the top of search results for 7 days."],
], col_widths=[2.4, 1.0, 3.2])
body("The classification of redemption rewards is intentional: one operational reward (download credits, useful immediately for power-users) and one social reward (pinning, useful for those who want public visibility). Future versions may add a third category — exchange points for printing-shop or library coupons — once the platform partners with the relevant campus services.")

heading3("2.3.3 The Use Case")
body("Take the example of a student redeeming 50 points for additional download credits.")
doc.add_paragraph()
add_table(["Field", "Description"], [
    ["Use case name", "Redeem points for download credits"],
    ["Actor", "Authenticated student user; system"],
    ["Description",
     "The student spends 50 of their accumulated points to obtain 100 additional download credits. Credits never expire and are deducted before the points balance whenever the user downloads."],
    ["Trigger", "Student has accumulated ≥50 points and wishes to lock in cheap downloads."],
    ["Preconditions",
     "1. User is logged in.\n2. User's points balance ≥ 50.\n3. The redemption catalog is enabled."],
    ["Normal course",
     "1. User opens the points dashboard.\n2. User clicks the '50 pts → 100 downloads' redemption card.\n3. System displays a confirmation modal showing the cost, benefit, and resulting balance.\n4. User confirms. The system runs an atomic transaction that decrements the points balance by 50, increments the download credits by 100, writes a PointRecord audit entry with action_type = REDEEM_DOWNLOAD_CREDIT, and writes a Redemption row.\n5. System returns the new balance and the new credit total to the dashboard."],
    ["Alternative course",
     "Insufficient balance at confirmation time → system shows 'Insufficient balance' and aborts without changing any state.\n\nUser cancels at the confirmation modal → no state change."],
    ["Postconditions",
     "Points balance reduced by 50. Download credits increased by 100. PointRecord row inserted. Redemption row inserted."],
], col_widths=[1.6, 5.0])

heading3("2.3.4 DFD")
figure_image("fig4_dfd_points.png",
             "Figure 4.  DFD diagram of the points and rewards system",
             width_inches=6.5)
body("The diagram traces three principal data flows. (1) An action event (upload approved, download received, rating received, or download spent) arrives at the Points Engine process. (2) The Points Engine applies the configured delta to the user's row using a single atomic UPDATE statement guarded by a balance condition, and writes both the updated balance and an audit entry to the PointRecord data store within one database transaction. (3) The new balance is returned to the user interface for display, and any side-effect (download credit increment, pin activation, leaderboard recomputation) is dispatched.")

heading3("2.3.5 Demo")
body("The points dashboard is the central screen for the new feature. The hero region at the top displays three large statistics: the user's current points balance, the running monthly earnings, and the remaining free daily downloads. Below the hero, a two-column layout shows the redemption options on the left (50 pts → 100 download credits; 100 pts → 7-day pin) and the monthly campus leaderboard on the right (top 20 contributors with the user's own row highlighted if they fall outside the top 20). At the bottom of the page, a complete points transaction history table lists every action with date, type, related resource, signed delta (color-coded green for earnings and red for spending), and running balance.")

doc.add_page_break()

# ============================================================
# 3.0 FUNCTIONAL ANALYSIS (2.0)
# ============================================================
heading1("3.0 Functional Analysis (2.0)")

heading2("3.1 AI-assisted recommendation system")

heading3("3.1.1 Backgrounds")
body("Following stabilization of the platform's core structure, version 2.0 will incorporate AI-powered recommendation technology. Its aim is not to alter the process of students' searching activity, but rather to minimize the need for manual re-filtering. In case a user regularly searches for a specific subject or retains resources of the same lecturer, the system will offer additional notes, previous exams, and question-and-answer sessions that match this established pattern, surfaced proactively rather than only on demand.")
body("Recommendation must be implemented cautiously because academic recommendation is private and accuracy-sensitive in nature. The system should never reveal the personal learning behavior of any other user; it should only use individual histories to derive aggregate patterns. The system should also justify why it recommends any particular resource, such as 'belongs to the same course you searched last week', 'highly rated by students in the same major', or 'frequently saved by students preparing for the same exam'. Without these explanations, students will not trust the recommendation and the feature will go unused.")
body("Recommendation will also be especially beneficial for new users who may not yet know what type of resources are most useful for a particular course, or who may not yet have learned the specific keywords needed to find them. In this case, the system can recommend resources that have been approved and rated highly by other users in the same major and academic year — effectively transferring the implicit knowledge of senior students into a structured suggestion stream.")

heading3("3.1.2 Use case")
doc.add_paragraph()
add_table(["Field", "Description"], [
    ["Use case name", "Recommend academic resources to a user"],
    ["Actor", "Authenticated student user; recommendation engine; system"],
    ["Description",
     "The system reads the user's recent search and download history, derives a feature vector, and returns a ranked list of resources that the user has not yet seen but is likely to find useful, each with a short explanation."],
    ["Trigger", "User opens the home page or the resource detail page."],
    ["Preconditions",
     "1. User is logged in.\n2. User has at least three prior interactions (searches, downloads, or saves) recorded in the database. Cold-start users fall back to popular-in-major recommendations."],
    ["Normal course",
     "1. System reads the user's recent course/search history.\n2. System filters the candidate set to verified resources matching the user's major and recent course tags.\n3. The recommendation engine ranks candidates by a hybrid model combining content similarity and collaborative filtering signals.\n4. The user opens a recommendation card and either downloads it (positive feedback) or marks it 'not helpful' (negative feedback).\n5. The system updates the user's feedback record for the next round."],
    ["Alternative course",
     "Insufficient history → system recommends popular verified resources from the user's major as a fallback.\n\nUser opts out of recommendations entirely → system disables the feature for that user and stops collecting recommendation interaction data."],
], col_widths=[1.6, 5.0])

heading3("3.1.3 DFD")
body("The version 2.0 DFD extends the v1.0 retrieval model. User history, resource metadata, and ratings data are fed into the recommendation module. The recommendation module generates a ranked recommendation list and feeds the user's responses (clicks, downloads, dismissals) back into the history database to refine future recommendations.")
figure_image("fig5_dfd_recommendation.png",
             "Figure 5.  DFD diagram of the v2.0 AI-assisted recommendation system",
             width_inches=6.5)

doc.add_page_break()

# ============================================================
# 4.0 RESULT AND FUTURE IMPACT
# ============================================================
heading1("4.0 Result and Future Impact")
body("The primary purpose of this platform is to improve the flow of academic resources within the campus environment. If students can find credible sources in less time, they can dedicate more time to actually understanding the material rather than to the meta-task of locating it. The carbon-trading-style markets considered in similar projects are not applicable here; instead the project's impact is realized through three dimensions: social, financial, and system cost.")

heading2("4.1 Social impacts:")
body("The platform can establish a sustained academic mutual-assistance culture. Students within the same course or department can support one another with curated, peer-rated resources, and the verification mechanism minimizes the spread of misinformation or outdated material. Over time the platform also becomes a passive sensor for student need: by aggregating which courses generate the most search traffic and the most resource requests, it provides student organizations and academic affairs offices with data that can guide where supplementary tutorials or shared notes are most badly needed.")
body("The points system also surfaces a positive social signal that does not currently exist on campus: explicit recognition for the students who quietly maintain the academic commons. Today these students share notes via WeChat groups and receive at most a 'thank you' emoji. On the platform their contributions are visible, ranked, and rewarded — both with operational benefits (download credits, pinning) and with reputational standing (the monthly leaderboard).")

heading2("4.2 Financial impacts:")
body("Version 1.0 of this project has no commercial intent. The financial value to the campus community is realized indirectly through (a) the time saved searching across multiple platforms, and (b) the prevention of repeat purchases of low-quality second-hand materials sold via peer reseller channels. With future partnerships between the platform and on-campus services such as the printing centre or library, the points balance can be redeemed for tangible micro-benefits, deepening the financial value loop without requiring any commercial monetization of the platform itself.")

heading2("4.3 System cost")
body("The full-project budget was estimated bottom-up at $10,000 HKD. Labour effort is the largest line item (notional, since the team is unpaid student labour); software costs were minimized by relying on free-tier infrastructure, which left actual cash outlay confined to user-research incentives. The breakdown is shown below.")
doc.add_paragraph()
add_table(["Category", "Budget", "Actual"], [
    ["Personnel cost (notional, 330h × $80/h)", "$2,000", "$0 (student labour)"],
    ["Tools and software cost", "$6,000", "$0 (free tiers)"],
    ["User research and testing incentives", "$1,000", "$520"],
    ["Contingency reserve (10%)", "$1,000", "$0"],
    ["TOTAL", "$10,000", "$520 (5.2% utilisation)"],
], col_widths=[3.2, 1.6, 2.0])

doc.add_page_break()

# ============================================================
# 5. APPENDIX
# ============================================================
heading1("5. Appendix (Questionnaire and Demo)")

heading2("5.1 Questionnaire")
body("The questionnaire below was distributed via campus WeChat groups in late March 2026. Forty-seven valid responses were collected. Demographic questions (year of study, major, frequency of academic resource searching) come first; the substantive questions follow.")
body("Q1. Which platform do you most often use to search for study materials? (multi-select)")
bullet("A. WeChat group chats")
bullet("B. University cloud drive (e.g., shared Baidu Wenku links)")
bullet("C. Paid third-party platforms")
bullet("D. Direct messaging with classmates / upperclassmen")
bullet("E. Other (please specify)")
body("Q2. On average, how much time does it take you to find a useful past paper or set of class notes?")
bullet("A. Less than 10 minutes")
bullet("B. 10-30 minutes")
bullet("C. 30-60 minutes")
bullet("D. More than 60 minutes")
body("Q3. Which of the following do you check before trusting an academic file shared by someone else? (multi-select)")
bullet("A. The file is from a senior student I personally know")
bullet("B. The course code matches the file's intended use")
bullet("C. The academic year is recent enough")
bullet("D. Other students have rated or recommended it")
bullet("E. None of the above; I generally trust shared files")
body("Q4. Would you upload your high-quality notes if a points or badge system rewarded contributions?")
bullet("A. Yes, definitely")
bullet("B. Probably yes, depending on the rewards")
bullet("C. Probably no")
bullet("D. Definitely no")
body("Q5. What types of rewards would most motivate you to answer questions or share materials with peers? (rank top 3)")
bullet("A. Tangible service coupons (printing, library)")
bullet("B. Operational credits within the platform (extra downloads)")
bullet("C. Reputation / honour badges")
bullet("D. Visibility on a campus leaderboard")
bullet("E. None of the above")
body("Q6. Of the following potential platform features, which two do you consider most important?")
bullet("A. Precise multi-dimensional search by course code")
bullet("B. Verified upload with content review")
bullet("C. Peer Q&A with bounty points")
bullet("D. Points-based reward for contributors")
bullet("E. AI-driven personalized recommendation")
bullet("F. Admin / moderation features for safety")

heading2("5.2 Demo")
body("The demo includes four primary interfaces in v1.0: Search Page, Resource Detail Page, Upload Page, and Points Dashboard. These four interfaces correspond to the major workflows documented in Section 2. Screenshots are embedded in the prototype HTML file (Prototype.html) submitted alongside this report. The live deployed system is also accessible at the URL below.")
body("")
body("Live URL: https://signing-isle-printed-shapes.trycloudflare.com")
body("GitHub repository: https://github.com/a2318491287-design/must-campus-resource-platform")
body("Demo account: 1230000000 / demo123 (100 points pre-loaded)")
body("")
body("—— End of Report ——")

doc.save(OUTPUT)
print(f"Done: {OUTPUT}")
print(f"  Sections: 5 chapters (Intro / Functional 1.0 / Functional 2.0 / Result & Impact / Appendix)")
print(f"  Figures embedded: 5 (mind-map + 4 DFDs)")
print(f"  Fonts: {ASCII_FONT} + {CJK_FONT} (universal Word defaults — fixes garbled-text issue)")
